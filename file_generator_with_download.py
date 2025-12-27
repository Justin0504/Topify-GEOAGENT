"""
required_open_webui_version: 0.6.0
description: Universal File Generator with Direct Download - Generate downloadable files that can be accessed directly from chat
requirements: reportlab, python-docx, openpyxl
"""

import base64
import io
import json
import csv
import os
import tempfile
import hashlib
from typing import Dict, Any, List, Optional
from datetime import datetime
from pydantic import BaseModel, Field

# PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_LEFT

# Word generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Excel generation
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill


class Tools:
    """Universal File Generator with Direct Download - 通用文件生成工具（支持直接下载）"""

    def __init__(self):
        # 使用 Open WebUI 的数据目录
        self.output_dir = os.path.join(os.getcwd(), "backend", "data", "generated_files")
        # 确保目录存在
        os.makedirs(self.output_dir, exist_ok=True)

    def _save_file_to_disk(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Save file to disk and return download information
        
        :param file_bytes: File content in bytes
        :param filename: Filename
        :return: File information including path and download URL
        """
        try:
            # 生成唯一文件名（添加时间戳和哈希避免冲突）
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            file_hash = hashlib.md5(file_bytes).hexdigest()[:8]
            name_parts = os.path.splitext(filename)
            unique_filename = f"{name_parts[0]}_{timestamp}_{file_hash}{name_parts[1]}"
            
            # 保存到磁盘
            file_path = os.path.join(self.output_dir, unique_filename)
            with open(file_path, 'wb') as f:
                f.write(file_bytes)
            
            # 生成下载信息
            download_url = f"/api/v1/files/{unique_filename}"
            
            return {
                "success": True,
                "filename": unique_filename,
                "original_filename": filename,
                "path": file_path,
                "size": len(file_bytes),
                "download_url": download_url,
                "download_link": f"[📥 下载 {filename}]({download_url})"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"保存文件失败: {str(e)}"
            }

    def _format_response(
        self,
        file_bytes: bytes,
        filename: str,
        file_format: str,
        include_base64: bool = False
    ) -> Dict[str, Any]:
        """
        Format response with file information
        
        :param file_bytes: File content
        :param filename: Filename
        :param file_format: File format (pdf, docx, etc.)
        :param include_base64: Whether to include Base64 data
        :return: Formatted response
        """
        # 保存文件到磁盘
        save_result = self._save_file_to_disk(file_bytes, filename)
        
        if not save_result["success"]:
            return save_result
        
        response = {
            "success": True,
            "filename": save_result["original_filename"],
            "saved_as": save_result["filename"],
            "format": file_format,
            "size": save_result["size"],
            "path": save_result["path"],
            "download_url": save_result["download_url"],
            "message": f"✅ 文件已生成: **{filename}** ({save_result['size']} bytes)\n\n{save_result['download_link']}\n\n💡 提示: 点击上面的链接直接下载文件"
        }
        
        # 可选：包含 Base64 数据（用于备用下载方式）
        if include_base64:
            base64_data = base64.b64encode(file_bytes).decode('utf-8')
            response["base64_data"] = base64_data
            response["download_command"] = f'echo "{base64_data}" | base64 -D > {filename} && open {filename}'
            response["message"] += f"\n\n📦 备用下载方式:\n在终端运行: `echo \"[base64_data]\" | base64 -D > {filename}`"
        
        return response

    def generate_pdf_document(
        self,
        title: str,
        content: str,
        filename: str = "document.pdf",
        author: str = "Open WebUI",
        page_size: str = "letter",
        include_base64: bool = False
    ) -> Dict[str, Any]:
        """
        Generate a PDF document with direct download link
        
        :param title: Document title
        :param content: Document content
        :param filename: Output filename
        :param author: Document author
        :param page_size: Page size ('letter' or 'A4')
        :param include_base64: Include Base64 data for alternative download
        :return: File data with download link
        """
        try:
            # Create PDF buffer
            buffer = io.BytesIO()
            
            # Select page size
            pagesize = A4 if page_size.lower() == "a4" else letter
            
            # Create PDF document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=pagesize,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18,
                title=title,
                author=author
            )
            
            # Define styles
            styles = getSampleStyleSheet()
            
            # Custom title style
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=30,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )
            
            # Custom body style
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=11,
                leading=16,
                spaceAfter=12,
                alignment=TA_LEFT
            )
            
            # Build content
            story = []
            
            # Add title
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.2 * inch))
            
            # Add metadata
            metadata = f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            metadata_style = ParagraphStyle(
                'Metadata',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.grey,
                alignment=TA_CENTER,
                spaceAfter=30
            )
            story.append(Paragraph(metadata, metadata_style))
            story.append(Spacer(1, 0.3 * inch))
            
            # Add content paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), body_style))
                    story.append(Spacer(1, 0.1 * inch))
            
            # Build PDF
            doc.build(story)
            
            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            # Format and return response
            return self._format_response(pdf_bytes, filename, "pdf", include_base64)
            
        except Exception as e:
            return {
                "success": False,
                "error": f"生成 PDF 文件失败: {str(e)}"
            }

    def generate_word_document(
        self,
        title: str,
        content: str,
        filename: str = "document.docx",
        author: str = "Open WebUI",
        include_base64: bool = False
    ) -> Dict[str, Any]:
        """
        Generate a Word document with direct download link
        
        :param title: Document title
        :param content: Document content
        :param filename: Output filename
        :param author: Document author
        :param include_base64: Include Base64 data for alternative download
        :return: File data with download link
        """
        try:
            # Create Word document
            doc = Document()
            
            # Set document properties
            doc.core_properties.author = author
            doc.core_properties.title = title
            doc.core_properties.created = datetime.now()
            
            # Add title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            metadata_para = doc.add_paragraph()
            metadata_para.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
            metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Spacer
            
            # Add content paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()
            
            # Format and return response
            return self._format_response(docx_bytes, filename, "docx", include_base64)
            
        except Exception as e:
            return {
                "success": False,
                "error": f"生成 Word 文件失败: {str(e)}"
            }

    def generate_excel_spreadsheet(
        self,
        data: List[List[Any]],
        filename: str = "spreadsheet.xlsx",
        sheet_name: str = "Sheet1",
        headers: Optional[List[str]] = None,
        title: Optional[str] = None,
        include_base64: bool = False
    ) -> Dict[str, Any]:
        """
        Generate an Excel spreadsheet with direct download link
        
        :param data: 2D list of data rows
        :param filename: Output filename
        :param sheet_name: Name of the worksheet
        :param headers: Optional list of column headers
        :param title: Optional spreadsheet title
        :param include_base64: Include Base64 data for alternative download
        :return: File data with download link
        """
        try:
            # Create workbook
            wb = Workbook()
            ws = wb.active
            ws.title = sheet_name
            
            current_row = 1
            
            # Add title if provided
            if title:
                ws.merge_cells(f'A1:{chr(65 + len(headers or data[0]) - 1)}1')
                title_cell = ws['A1']
                title_cell.value = title
                title_cell.font = Font(size=16, bold=True, color="FFFFFF")
                title_cell.fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
                title_cell.alignment = Alignment(horizontal="center", vertical="center")
                ws.row_dimensions[1].height = 30
                current_row = 2
            
            # Add headers if provided
            if headers:
                for col_idx, header in enumerate(headers, start=1):
                    cell = ws.cell(row=current_row, column=col_idx)
                    cell.value = header
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill(start_color="3498db", end_color="3498db", fill_type="solid")
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                current_row += 1
            
            # Add data
            for row_data in data:
                for col_idx, value in enumerate(row_data, start=1):
                    ws.cell(row=current_row, column=col_idx, value=value)
                current_row += 1
            
            # Auto-adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = None
                for cell in column:
                    try:
                        # Skip merged cells
                        if hasattr(cell, 'column_letter'):
                            if column_letter is None:
                                column_letter = cell.column_letter
                            if cell.value and len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                    except:
                        pass
                if column_letter:
                    adjusted_width = min(max_length + 2, 50)
                    ws.column_dimensions[column_letter].width = adjusted_width
            
            # Save to buffer
            buffer = io.BytesIO()
            wb.save(buffer)
            xlsx_bytes = buffer.getvalue()
            buffer.close()
            
            # Format and return response
            return self._format_response(xlsx_bytes, filename, "xlsx", include_base64)
            
        except Exception as e:
            return {
                "success": False,
                "error": f"生成 Excel 文件失败: {str(e)}"
            }

    def generate_text_file(
        self,
        content: str,
        filename: str = "document.txt",
        encoding: str = "utf-8",
        include_base64: bool = False
    ) -> Dict[str, Any]:
        """
        Generate a plain text file with direct download link
        
        :param content: Text content
        :param filename: Output filename
        :param encoding: Text encoding
        :param include_base64: Include Base64 data for alternative download
        :return: File data with download link
        """
        try:
            text_bytes = content.encode(encoding)
            return self._format_response(text_bytes, filename, "text", include_base64)
        except Exception as e:
            return {
                "success": False,
                "error": f"生成文本文件失败: {str(e)}"
            }

    def quick_generate_file(
        self,
        content: str,
        file_type: str = "pdf",
        title: str = "Document",
        filename: Optional[str] = None,
        include_base64: bool = False
    ) -> Dict[str, Any]:
        """
        Quick file generation with automatic format detection and direct download
        
        :param content: File content
        :param file_type: File type (pdf, docx, txt)
        :param title: Document title
        :param filename: Optional custom filename
        :param include_base64: Include Base64 for alternative download
        :return: File data with download link
        """
        file_type = file_type.lower()
        
        # Auto-generate filename if not provided
        if not filename:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            extensions = {
                'pdf': 'pdf',
                'docx': 'docx',
                'word': 'docx',
                'txt': 'txt',
                'text': 'txt',
            }
            ext = extensions.get(file_type, 'txt')
            filename = f"{title}_{timestamp}.{ext}"
        
        # Route to appropriate generator
        if file_type in ['pdf']:
            return self.generate_pdf_document(title, content, filename, include_base64=include_base64)
        elif file_type in ['docx', 'word']:
            return self.generate_word_document(title, content, filename, include_base64=include_base64)
        elif file_type in ['txt', 'text']:
            return self.generate_text_file(content, filename, include_base64=include_base64)
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {file_type}. Supported: pdf, docx, txt"
            }

