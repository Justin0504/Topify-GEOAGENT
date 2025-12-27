"""
title: 文章写作工具
description: 【单篇文章生成】根据搜索意图写作GEO优化文章 | 【完整文章生成】倒金字塔结构完整文章（含配图、链接） | 【批量文章生成】为多个主题批量生成文章
author: GEO Agent
version: 2.0.0
required_open_webui_version: 0.6.0
requirements: python-docx, requests, beautifulsoup4
"""

import os
import re
import requests
import urllib3
from typing import List, Optional
from datetime import datetime
from pydantic import BaseModel, Field
from urllib.parse import urlparse, urljoin
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 禁用SSL警告
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)


class Tools:
    """
    文章写作工具 - 生成SEO/GEO优化的文章内容
    
    ═══════════════════════════════════════════════════════════════
    🎯 功能匹配指南（中文触发词）
    ═══════════════════════════════════════════════════════════════
    
    ✍️ write_single_article - 单篇文章框架生成
       触发词: "写文章框架", "生成文章大纲", "创建文章结构",
              "文章框架", "GEO文章框架"
       示例: "为关键词 best AI tools 生成文章框架"
       输出: Word 文档（包含GEO优化结构框架）
    
    📝 write_complete_article - 完整文章生成（支持WordPress发布）
       触发词: "写完整文章", "生成完整文章", "倒金字塔文章",
              "分析搜索意图并写文章", "写一篇完整的文章",
              "写文章并发布", "发布到WordPress"
       使用流程:
        1. LLM先根据搜索关键词、产品信息生成完整的文章内容（Markdown格式）
        2. 然后调用此工具，传入生成的 article_content
        3. 工具会自动生成SEO友好的AI图片URL、添加链接、生成Word文档、可选发布到WordPress
       示例: 
        - "分析搜索主题「best AI SEO tools」的搜索意图，然后写一篇满足用户搜索意图的文章，推荐 Topify.ai" 
          → LLM生成内容 → 调用工具（仅生成Word）
        - "写文章推荐Topify并发布到WordPress" 
          → LLM生成内容 → 调用工具（生成Word并发布）
       输出: Word 文档（完整的倒金字塔结构文章，含配图、链接）+ 可选WordPress发布
       注意: ⚠️ article_content 参数必填，工具不会自动生成内容
    
    📚 write_batch_articles - 批量文章生成
       触发词: "批量文章", "多篇文章", "30篇文章", "批量写作",
              "批量生成", "多个主题"
       示例: "为以下30个主题生成文章"
       输出: 多个Word文档或汇总文档
    
    ═══════════════════════════════════════════════════════════════
    """

    class Valves(BaseModel):
        OUTPUT_PATH: str = Field(
            default="/app/backend/data/output",
            description="文件保存路径（Docker环境）"
        )
        DEFAULT_LANGUAGE: str = Field(
            default="en",
            description="默认文章语言 (en=英文, zh=中文)"
        )
        WP_ACCESS_TOKEN: str = Field(
            default="",
            description="【可选】WordPress.com API Access Token（用于自动发布）"
        )
        WP_SITE_ID: str = Field(
            default="",
            description="【可选】WordPress.com Site ID（用于自动发布）"
        )
        WP_API_BASE: str = Field(
            default="https://public-api.wordpress.com/rest/v1.1",
            description="WordPress.com API 基础 URL"
        )

    def __init__(self):
        self.valves = self.Valves()

    def _aw_ensure_output_dir(self) -> str:
        """确保输出目录存在"""
        output_path = self.valves.OUTPUT_PATH
        if not os.path.exists(output_path):
            os.makedirs(output_path, exist_ok=True)
        return output_path

    def _create_article_doc(
        self,
        title: str,
        content_sections: List,
        product_name: str,
        product_url: str,
        faqs: List = None
    ):
        """创建文章Word文档"""
        doc = Document()
        
        # 标题
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # TL;DR 摘要
        tldr = doc.add_paragraph()
        tldr_run = tldr.add_run('TL;DR: ')
        tldr_run.bold = True
        tldr.add_run(content_sections[0].get('tldr', '[摘要内容]'))
        
        doc.add_paragraph()  # 空行
        
        # 正文内容
        for section in content_sections:
            if section.get('heading'):
                doc.add_heading(section['heading'], level=2)
            if section.get('content'):
                doc.add_paragraph(section['content'])
        
        # 产品推荐部分
        doc.add_heading(f'Why Choose {product_name}?', level=2)
        rec_para = doc.add_paragraph()
        rec_para.add_run(f'Based on our analysis, ')
        link_run = rec_para.add_run(product_name)
        link_run.bold = True
        link_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
        rec_para.add_run(f' stands out as a top choice. ')
        rec_para.add_run(f'Learn more at: {product_url}')
        
        # FAQ 部分
        if faqs:
            doc.add_heading('Frequently Asked Questions', level=2)
            for faq in faqs:
                q_para = doc.add_paragraph()
                q_run = q_para.add_run(f"Q: {faq.get('question', '')}")
                q_run.bold = True
                
                a_para = doc.add_paragraph()
                a_para.add_run(f"A: {faq.get('answer', '')}")
                doc.add_paragraph()  # 空行
        
        return doc

    def write_single_article(
        self,
        keyword: str,
        product_name: str,
        product_url: str,
        product_description: str,
        search_intent: str = "",
        word_count: int = 1500,
        language: str = "en"
    ) -> str:
        """
        【单篇文章写作工具】根据搜索关键词和意图写作GEO优化文章
        
        当用户说以下内容时调用此工具：
        - "写文章"、"生成文章"、"创建文章"
        - "写一篇关于xxx的文章"
        - "为关键词xxx写文章"
        - "GEO优化文章"
        
        :param keyword: 目标搜索关键词
        :param product_name: 要推荐的产品名称
        :param product_url: 产品官网URL
        :param product_description: 产品描述
        :param search_intent: 搜索意图分析（可选）
        :param word_count: 目标字数
        :param language: 文章语言 (en=英文, zh=中文)
        :return: 包含Word文件路径的结果，以及文章的完整内容
        """
        output_path = self._aw_ensure_output_dir()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_keyword = keyword.replace(' ', '_').replace('/', '_')[:30]
        filename = f"article_{safe_keyword}_{timestamp}.docx"
        filepath = os.path.join(output_path, filename)
        
        # 生成文章结构（这里提供框架，实际内容由LLM生成）
        article_structure = f"""
# {keyword.title()}

**TL;DR**: [在此处用2-3句话直接回答"{keyword}"这个问题，包含关键数据点，并提及{product_name}作为推荐方案。]

## What is {keyword.split()[0] if keyword.split() else keyword}?

[开头段落：在前60个词内直接回答核心问题。使用具体数据和事实。]

[解释段落：详细说明概念、背景和重要性。]

## Key Features to Look For

When evaluating solutions for {keyword}, consider these essential factors:

1. **[特性1]** - [说明]
2. **[特性2]** - [说明]
3. **[特性3]** - [说明]
4. **[特性4]** - [说明]

## Top Solutions Comparison

| Solution | Key Strength | Best For | Rating |
|----------|-------------|----------|--------|
| {product_name} | [核心优势] | [目标用户] | ⭐⭐⭐⭐⭐ |
| [竞品1] | [优势] | [用户] | ⭐⭐⭐⭐ |
| [竞品2] | [优势] | [用户] | ⭐⭐⭐ |

## Why {product_name} Stands Out

{product_description}

Key advantages of {product_name}:
- [优势1]
- [优势2]
- [优势3]

👉 **Learn more**: [{product_name}]({product_url})

## How to Get Started

Step-by-step guide:

1. **Step 1**: [具体操作]
2. **Step 2**: [具体操作]
3. **Step 3**: [具体操作]

## Real-World Use Cases

### Use Case 1: [场景名称]
[具体案例描述]

### Use Case 2: [场景名称]
[具体案例描述]

## Frequently Asked Questions

### Q: [常见问题1]?
A: [详细回答，自然地提及{product_name}的相关功能]

### Q: [常见问题2]?
A: [详细回答]

### Q: [常见问题3]?
A: [详细回答]

## Conclusion

[总结段落：重申核心观点，强调{product_name}的价值，包含行动号召。]

**Ready to get started?** Visit [{product_name}]({product_url}) today.

---
*Last updated: {datetime.now().strftime("%B %Y")}*
"""

        # 创建Word文档
        doc = Document()
        
        # 标题
        doc.add_heading(keyword.title(), 0)
        
        # 元信息
        meta = doc.add_paragraph()
        meta.add_run(f'Target Keyword: ').bold = True
        meta.add_run(keyword)
        meta.add_run(f'\nProduct: ').bold = True
        meta.add_run(product_name)
        meta.add_run(f'\nWord Count Target: ').bold = True
        meta.add_run(f'{word_count} words')
        meta.add_run(f'\nLanguage: ').bold = True
        meta.add_run('English' if language == 'en' else '中文')
        
        doc.add_paragraph()
        
        # 搜索意图分析
        doc.add_heading('Search Intent Analysis', level=1)
        if search_intent:
            doc.add_paragraph(search_intent)
        else:
            doc.add_paragraph(f'''
Based on the keyword "{keyword}", the search intent appears to be:
- Intent Type: [Informational/Commercial/Transactional]
- User Goal: [用户想要解决什么问题]
- Content Angle: [应该从什么角度写作]
''')
        
        # 文章大纲
        doc.add_heading('Article Outline', level=1)
        doc.add_paragraph(article_structure)
        
        # 写作指南
        doc.add_page_break()
        doc.add_heading('Writing Guidelines (GEO Optimized)', level=1)
        
        guidelines = [
            ('开头', '在前40-60个词内直接回答核心问题，便于AI引用'),
            ('TL;DR', '提供2-3句话摘要，包含关键数据和推荐'),
            ('结构', '使用清晰的H2/H3层级，每段2-4句'),
            ('数据', '每150-200词包含一个具体数据/统计'),
            ('FAQ', '添加3-5个常见问题，用问答形式'),
            ('产品提及', f'自然地在2-3处提及{product_name}'),
            ('链接', f'在合适位置添加{product_url}链接'),
            ('号召行动', '结尾包含明确的CTA')
        ]
        
        for item, desc in guidelines:
            p = doc.add_paragraph()
            p.add_run(f'• {item}: ').bold = True
            p.add_run(desc)
        
        doc.save(filepath)
        
        return f"""
📝 **文章框架生成完成**

🎯 目标关键词: {keyword}
📦 推荐产品: {product_name}
🔗 产品链接: {product_url}
📊 目标字数: {word_count} 词
🌐 语言: {'英文' if language == 'en' else '中文'}

═══════════════════════════════════════
📋 **文章结构**
═══════════════════════════════════════
1. TL;DR 摘要
2. 概念介绍
3. 关键特性
4. 方案对比（含产品推荐）
5. 产品优势
6. 使用指南
7. 实际案例
8. FAQ
9. 总结与CTA

═══════════════════════════════════════
💾 **文件已保存**
═══════════════════════════════════════
路径: {filepath}

═══════════════════════════════════════
📄 **文章大纲预览**
═══════════════════════════════════════
{article_structure[:1500]}...

💡 **下一步**:
请基于以上大纲完成文章正文写作，确保：
- 开头直接回答问题
- 自然融入产品推荐
- 包含具体数据和案例
- 添加FAQ部分
"""

    def write_batch_articles(
        self,
        topics: str,
        product_name: str,
        product_url: str,
        product_description: str
    ) -> str:
        """
        【批量文章写作工具】为多个主题批量生成文章框架
        
        当用户说以下内容时调用此工具：
        - "批量文章"、"多篇文章"
        - "为以下主题生成文章"
        - "批量写作30篇"
        - "多个关键词写文章"
        
        :param topics: 文章主题列表（每行一个主题，或用逗号分隔）
        :param product_name: 要推荐的产品名称
        :param product_url: 产品官网URL
        :param product_description: 产品描述
        :return: 包含批量生成结果的报告
        """
        output_path = self._aw_ensure_output_dir()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # 解析主题列表
        topic_list = []
        for line in topics.replace(',', '\n').split('\n'):
            topic = line.strip()
            if topic and not topic.startswith('#'):
                topic_list.append(topic)
        
        if not topic_list:
            return "❌ 未检测到有效的文章主题，请提供主题列表（每行一个或用逗号分隔）"
        
        # 创建汇总文档
        summary_filename = f"batch_articles_summary_{timestamp}.docx"
        summary_filepath = os.path.join(output_path, summary_filename)
        
        doc = Document()
        doc.add_heading('批量文章写作计划', 0)
        
        # 概述
        doc.add_heading('项目概述', level=1)
        overview = doc.add_paragraph()
        overview.add_run(f'产品: ').bold = True
        overview.add_run(f'{product_name}\n')
        overview.add_run(f'网站: ').bold = True
        overview.add_run(f'{product_url}\n')
        overview.add_run(f'文章数量: ').bold = True
        overview.add_run(f'{len(topic_list)} 篇\n')
        overview.add_run(f'生成时间: ').bold = True
        overview.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
        
        # 文章列表
        doc.add_heading('文章主题列表', level=1)
        
        article_summaries = []
        
        for idx, topic in enumerate(topic_list, 1):
            doc.add_heading(f'{idx}. {topic}', level=2)
            
            # 搜索意图分析提示
            intent_para = doc.add_paragraph()
            intent_para.add_run('搜索意图: ').bold = True
            intent_para.add_run('[待分析 - Informational/Commercial/Transactional]')
            
            # 文章结构
            structure_para = doc.add_paragraph()
            structure_para.add_run('建议结构:\n').bold = True
            structure_para.add_run(f'''
• TL;DR: 直接回答"{topic}"的核心问题
• 介绍: 什么是{topic.split()[0] if topic.split() else topic}
• 对比: 列出3-5个方案，突出{product_name}
• 指南: 如何使用/选择
• FAQ: 3个常见问题
• 结论: 推荐{product_name}，包含CTA
''')
            
            article_summaries.append({
                'id': idx,
                'topic': topic,
                'status': '待写作'
            })
            
            doc.add_paragraph()  # 空行分隔
        
        # 写作指南
        doc.add_page_break()
        doc.add_heading('GEO优化写作指南', level=1)
        
        guidelines = [
            '1. 每篇文章在开头40-60词内直接回答核心问题',
            '2. 添加TL;DR摘要，便于AI快速提取信息',
            '3. 使用清晰的H2/H3层级结构',
            '4. 每150-200词包含一个具体数据或统计',
            '5. 添加FAQ部分（3-5个问答）',
            f'6. 自然地在文章中2-3处提及{product_name}',
            '7. 确保产品链接可点击',
            '8. 结尾包含明确的行动号召（CTA）'
        ]
        
        for guideline in guidelines:
            doc.add_paragraph(guideline, style='List Bullet')
        
        # 产品信息
        doc.add_heading('产品信息（写作参考）', level=1)
        doc.add_paragraph(f'产品名称: {product_name}')
        doc.add_paragraph(f'官网地址: {product_url}')
        doc.add_paragraph(f'产品描述: {product_description}')
        
        doc.save(summary_filepath)
        
        # 生成文章列表表格
        articles_table = "\n".join([
            f"| {a['id']} | {a['topic'][:40]}{'...' if len(a['topic']) > 40 else ''} | {a['status']} |"
            for a in article_summaries[:20]  # 只显示前20个
        ])
        
        return f"""
📚 **批量文章计划生成完成**

📦 产品: {product_name}
🔗 链接: {product_url}
📄 文章数量: {len(topic_list)} 篇

═══════════════════════════════════════
📋 **文章列表预览**
═══════════════════════════════════════
| # | 主题 | 状态 |
|---|------|------|
{articles_table}
{f'... 还有 {len(topic_list) - 20} 篇' if len(topic_list) > 20 else ''}

═══════════════════════════════════════
💾 **文件已保存**
═══════════════════════════════════════
路径: {summary_filepath}

📋 文档包含:
1. 项目概述
2. {len(topic_list)}个文章主题及写作建议
3. GEO优化写作指南
4. 产品信息参考

═══════════════════════════════════════
💡 **下一步建议**
═══════════════════════════════════════
1. 逐一为每个主题写作完整文章
2. 使用 write_single_article 工具生成单篇详细框架
3. 完成后使用 WordPress 工具批量上传
4. 优化TDK后批量发布

🤖 **批量写作提示**:
你可以说"为第1-10个主题写文章"来开始批量写作
"""

    def _aw_extract_images_from_url(self, url: str, max_images: int = 5) -> list:
        """从网页URL中提取高质量图片URL列表（支持静态和动态加载的图片）"""
        try:
            from bs4 import BeautifulSoup
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(url, timeout=15, verify=False, allow_redirects=True, headers=headers)
            if response.status_code == 200:
                soup = BeautifulSoup(response.content, 'html.parser')
                html_content = response.text  # 保存原始HTML用于正则提取
                image_candidates = []  # 存储候选图片（URL, 优先级分数）
                
                # 需要过滤的关键词（在URL或class/id中）
                exclude_keywords = ['logo', 'icon', 'avatar', 'favicon', 'button', 'badge', 
                                   'spinner', 'loader', 'placeholder', 'thumbnail', 
                                   'screenshot', 'screen-shot', 'mockup', 'frame', 'browser',
                                   'window', 'desktop', 'capture', 'preview-img', 'dashboard',
                                   'chart', 'graph', 'analytics', 'metrics', 'ui-', '-ui']
                
                # 优先查找的class/id关键词（hero图片、banner等）
                priority_keywords = ['hero', 'banner', 'feature', 'main', 'primary', 
                                    'showcase', 'demo', 'product', 'illustration']
                
                # 提取所有 img 标签（扩大搜索范围）
                for img in soup.find_all('img'):
                    src = (img.get('src') or img.get('data-src') or 
                          img.get('data-lazy-src') or img.get('data-original'))
                    if not src:
                        continue
                    
                    # 处理相对URL
                    if src.startswith('//'):
                        src = 'https:' + src
                    elif src.startswith('/') or not src.startswith('http'):
                        src = urljoin(url, src)
                    
                    # 过滤SVG（通常是图标）
                    if src.lower().endswith('.svg'):
                        continue
                    
                    # 检查URL中是否包含排除关键词
                    src_lower = src.lower()
                    # 严格检查是否是logo
                    is_logo_src = any(pattern in src_lower for pattern in ['/logo', 'logo/', 'logo.', '-logo.', '.logo', 'favicon'])
                    # 检查是否是截图或仪表板
                    is_screenshot_src = any(pattern in src_lower for pattern in ['screenshot', 'screen-shot', 'mockup', 'dashboard', 'browser-', 'window-', 'chart', 'graph', 'analytics', 'metrics'])
                    
                    if is_logo_src or is_screenshot_src or any(keyword in src_lower for keyword in exclude_keywords):
                        continue
                    
                    # 检查class和id中是否包含排除关键词
                    img_class = img.get('class', [])
                    img_id = img.get('id', '')
                    img_classes_str = ' '.join(img_class).lower() + ' ' + img_id.lower()
                    if any(keyword in img_classes_str for keyword in exclude_keywords):
                        continue
                    
                    # 计算优先级分数
                    priority = 0
                    
                    # 优先级：尺寸大的图片得分高
                    width = img.get('width') or img.get('data-width')
                    height = img.get('height') or img.get('data-height')
                    has_size_info = False
                    if width and height:
                        try:
                            w, h = int(width), int(height)
                            has_size_info = True
                            # 降低尺寸要求：从200px降到150px，如果还是太小，再降到100px
                            min_size = 150
                            if w < min_size or h < min_size:
                                # 如果尺寸太小，降低优先级但不完全排除
                                if w < 100 or h < 100:
                                    continue  # 小于100px的完全排除
                                priority += 10  # 小图片给低优先级
                            else:
                                # 尺寸越大，优先级越高
                                priority += min(w * h / 10000, 100)  # 最大100分
                        except (ValueError, TypeError):
                            has_size_info = False
                    
                    if not has_size_info:
                        # 没有尺寸信息，给中等优先级（不排除，因为很多现代网站使用CSS控制尺寸）
                        priority += 40  # 给中等优先级，允许没有尺寸信息的图片
                    
                    # 优先级：包含优先关键词的图片得分更高
                    if any(keyword in img_classes_str for keyword in priority_keywords):
                        priority += 50
                    if any(keyword in src_lower for keyword in priority_keywords):
                        priority += 30
                    
                    # 优先级：alt文本中有意义的描述（不是空或单个词）
                    alt = img.get('alt', '')
                    if alt and len(alt.split()) > 1:
                        priority += 20
                    
                    # 只接受常见图片格式，并去重
                    if any(src_lower.endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp']):
                        # 去重：检查基础URL
                        base_url = src.split('?')[0]
                        if base_url not in [url.split('?')[0] for url, _ in image_candidates]:
                            image_candidates.append((src, priority))
                
                # ===== 方法2：从CSS background-image提取（处理动态加载的图片）=====
                # 查找所有元素的 style 属性中的 background-image
                for element in soup.find_all(style=True):
                    style = element.get('style', '')
                    # 匹配 background-image: url(...)
                    bg_matches = re.findall(r'background-image\s*:\s*url\(["\']?([^"\')]+)["\']?\)', style)
                    for bg_url in bg_matches:
                        if bg_url.startswith('//'):
                            bg_url = 'https:' + bg_url
                        elif bg_url.startswith('/') or not bg_url.startswith('http'):
                            bg_url = urljoin(url, bg_url)
                        # 过滤明显的小图标、logo和截图
                        bg_url_lower = bg_url.lower()
                        is_logo_bg = any(pattern in bg_url_lower for pattern in ['/logo', 'logo/', 'logo.', '-logo.', '.logo', 'favicon'])
                        is_screenshot_bg = any(pattern in bg_url_lower for pattern in ['screenshot', 'screen-shot', 'mockup', 'dashboard', 'browser-', 'window-'])
                        
                        if not is_logo_bg and not is_screenshot_bg:
                            exclude_bg = ['icon', 'favicon', 'avatar', 'screenshot', 'mockup', 'frame', 'browser', 'dashboard', 'chart', 'graph']
                            if not any(exclude in bg_url_lower for exclude in exclude_bg):
                                if any(bg_url_lower.endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp']):
                                    # background-image 的图片通常是大图片，给较高优先级
                                    base_url = bg_url.split('?')[0]
                                    if base_url not in [url.split('?')[0] for url, _ in image_candidates]:
                                        image_candidates.append((bg_url, 60))
                
                # ===== 方法3：从HTML源代码中提取图片URL（包括JavaScript中的）=====
                # 查找常见的图片URL模式（http/https开头的图片URL）
                # 优先匹配完整的URL（带协议）
                full_url_pattern = r'https?://[^\s"\'<>\)]+\.(?:jpg|jpeg|png|gif|webp)(?:\?[^\s"\'<>\)]*)?'
                full_url_matches = re.findall(full_url_pattern, html_content, re.IGNORECASE)
                for img_url in full_url_matches:
                    # 清理URL：去除HTML实体编码，清理尾部标点
                    img_url = img_url.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
                    img_url = img_url.strip().rstrip('.,;:!?)')
                    # 如果是srcSet格式，提取第一个URL
                    if ' ' in img_url and img_url.find('.jpg') > 0:
                        img_url = img_url.split()[0]
                    
                    src_lower = img_url.lower()
                    # CDN图片通常质量较高，不过滤（如framerusercontent.com, cloudinary.com等）
                    is_cdn_image = any(cdn in src_lower for cdn in ['framerusercontent.com', 'cloudinary.com', 'cdn.', 'images.', 'assets.', 'imgix.net'])
                    # 排除明显的logo、图标、截图、仪表板
                    exclude_patterns = ['favicon', '-logo', 'logo', '/logo', 'logo/', 'icon-', '-icon', 
                                      'screenshot', 'screen-shot', 'mockup', 'dashboard', 'chart', 'graph',
                                      'analytics', 'metrics', 'browser-', '-browser', 'window-', '-window']
                    
                    # 检查是否是logo（通常在文件名或路径中）
                    is_logo = any(pattern in src_lower for pattern in ['/logo', 'logo/', 'logo.', '-logo.', '.logo', 'favicon'])
                    # 检查是否是截图或仪表板
                    is_screenshot = any(pattern in src_lower for pattern in ['screenshot', 'screen-shot', 'mockup', 'dashboard', 'browser-', 'window-'])
                    
                    # 如果是CDN图片，但明显是logo或截图，也要过滤
                    if is_logo or is_screenshot:
                        continue
                    
                    if (not any(exclude in src_lower for exclude in exclude_patterns) or (is_cdn_image and not is_logo and not is_screenshot)):
                        # 确保是图片格式
                        if any(ext in src_lower.split('?')[0] for ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp']):
                            # 去重：检查基础URL（去除查询参数）是否已存在
                            base_url = img_url.split('?')[0]
                            if base_url not in [url.split('?')[0] for url, _ in image_candidates]:
                                # CDN图片给较高优先级
                                priority = 55 if is_cdn_image else 35
                                image_candidates.append((img_url, priority))
                
                # 也匹配引号中的相对路径或URL
                quoted_pattern = r'["\']([^"\']*\.(?:jpg|jpeg|png|gif|webp)(?:\?[^"\']*)?)["\']'
                quoted_matches = re.findall(quoted_pattern, html_content, re.IGNORECASE)
                for match in quoted_matches:
                    img_url = match.strip().rstrip('.,;:!?)')
                    if img_url.startswith('//'):
                        img_url = 'https:' + img_url
                    elif img_url.startswith('/') or not img_url.startswith('http'):
                        img_url = urljoin(url, img_url)
                    src_lower = img_url.lower()
                    is_cdn_image = any(cdn in src_lower for cdn in ['framerusercontent.com', 'cloudinary.com', 'cdn.', 'images.', 'assets.', 'imgix.net'])
                    exclude_patterns = ['favicon', '-logo', 'logo', '/logo', 'logo/', 'icon-', '-icon', 
                                      'screenshot', 'screen-shot', 'mockup', 'dashboard', 'chart', 'graph',
                                      'analytics', 'metrics', 'browser-', '-browser', 'window-', '-window']
                    
                    # 检查是否是logo或截图
                    is_logo = any(pattern in src_lower for pattern in ['/logo', 'logo/', 'logo.', '-logo.', '.logo', 'favicon'])
                    is_screenshot = any(pattern in src_lower for pattern in ['screenshot', 'screen-shot', 'mockup', 'dashboard', 'browser-', 'window-'])
                    
                    if is_logo or is_screenshot:
                        continue
                    
                    if (not any(exclude in src_lower for exclude in exclude_patterns) or (is_cdn_image and not is_logo and not is_screenshot)):
                        if any(ext in src_lower.split('?')[0] for ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp']):
                            # 去重：检查基础URL
                            base_url = img_url.split('?')[0]
                            if base_url not in [url.split('?')[0] for url, _ in image_candidates]:
                                priority = 55 if is_cdn_image else 35
                                image_candidates.append((img_url, priority))
                
                # ===== 方法4：从JSON-LD结构化数据中提取 =====
                json_ld_scripts = soup.find_all('script', type='application/ld+json')
                for script in json_ld_scripts:
                    try:
                        import json
                        data = json.loads(script.string)
                        # 递归查找图片URL
                        def find_images_in_json(obj, skip_logo=False):
                            images = []
                            if isinstance(obj, dict):
                                for key, value in obj.items():
                                    # 跳过logo字段
                                    is_logo_field = key.lower() in ['logo', 'logourl', 'logourl']
                                    if key in ['image', 'photo', 'picture', 'thumbnail', 'thumbnailUrl'] or (not skip_logo and is_logo_field):
                                        if isinstance(value, str) and any(value.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp']):
                                            # 检查是否是logo
                                            value_lower = value.lower()
                                            is_logo_img = any(pattern in value_lower for pattern in ['/logo', 'logo/', 'logo.', '-logo.', '.logo', 'favicon'])
                                            if not is_logo_img or skip_logo:
                                                images.append((value, is_logo_field))
                                    else:
                                        images.extend(find_images_in_json(value, skip_logo))
                            elif isinstance(obj, list):
                                for item in obj:
                                    images.extend(find_images_in_json(item, skip_logo))
                            return images
                        json_images = find_images_in_json(data)
                        for json_img, is_logo in json_images:
                            if is_logo:  # 跳过logo
                                continue
                            if json_img.startswith('//'):
                                json_img = 'https:' + json_img
                            elif json_img.startswith('/') or not json_img.startswith('http'):
                                json_img = urljoin(url, json_img)
                            
                            # 再次检查是否是logo或截图
                            json_img_lower = json_img.lower()
                            is_logo_check = any(pattern in json_img_lower for pattern in ['/logo', 'logo/', 'logo.', '-logo.', '.logo', 'favicon'])
                            is_screenshot_check = any(pattern in json_img_lower for pattern in ['screenshot', 'screen-shot', 'mockup', 'dashboard', 'browser-', 'window-'])
                            if is_logo_check or is_screenshot_check:
                                continue
                            
                            # 去重
                            base_url = json_img.split('?')[0]
                            if base_url not in [url.split('?')[0] for url, _ in image_candidates]:
                                image_candidates.append((json_img, 70))  # JSON-LD中的图片通常质量较高
                    except (json.JSONDecodeError, ValueError, AttributeError):
                        pass
                
                # 按优先级排序，取前max_images个
                image_candidates.sort(key=lambda x: x[1], reverse=True)
                
                # 去重：基于基础URL（去除查询参数）避免重复图片
                seen_base_urls = set()
                unique_images = []
                for img_url, priority in image_candidates:
                    # 提取基础URL（去除查询参数）
                    base_url = img_url.split('?')[0]
                    if base_url not in seen_base_urls:
                        seen_base_urls.add(base_url)
                        unique_images.append(img_url)
                        if len(unique_images) >= max_images:
                            break
                
                return unique_images
        except Exception as e:
            # 添加调试信息（可选，可以通过日志查看）
            import logging
            logging.debug(f"图片提取失败: {str(e)}")
            pass
        return []
    
    def _aw_download_image(self, image_url: str, timeout: int = 10) -> Optional[BytesIO]:
        """从URL下载图片"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                'Accept': 'image/webp,image/apng,image/*,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.9',
            }
            response = requests.get(image_url, timeout=timeout, verify=False, allow_redirects=True, headers=headers)
            if response.status_code == 200:
                content_type = response.headers.get('content-type', '').lower()
                # 检查是否是图片类型
                if content_type.startswith('image/'):
                    return BytesIO(response.content)
                # 如果Content-Type不明确，检查文件扩展名
                elif any(image_url.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp']):
                    return BytesIO(response.content)
        except Exception as e:
            pass
        return None
    
    def _aw_generate_seo_image_url(self, keyword: str, image_index: int = 1, base_domain: str = None) -> str:
        """生成SEO友好的AI生成图片URL"""
        # 清理关键词，生成URL友好的slug
        slug = keyword.lower().strip()
        # 替换空格为连字符
        slug = re.sub(r'\s+', '-', slug)
        # 移除特殊字符，只保留字母、数字、连字符
        slug = re.sub(r'[^a-z0-9\-]', '', slug)
        # 移除多余的连字符
        slug = re.sub(r'-+', '-', slug)
        slug = slug.strip('-')
        # 限制长度
        slug = slug[:50]
        
        # 生成图片文件名（SEO友好）
        image_filename = f"{slug}-ai-generated-{image_index}.jpg"
        
        # 如果有base_domain，使用它；否则使用占位符URL
        if base_domain:
            # 从URL中提取域名
            domain = base_domain.replace('https://', '').replace('http://', '').split('/')[0]
            return f"https://{domain}/images/{image_filename}"
        else:
            # 使用占位符URL格式（实际使用时需要替换为真实的图片服务URL）
            return f"https://images.example.com/{image_filename}"
    
    def _aw_format_cell_text(self, text: str, product_name: str, product_url: str) -> str:
        """格式化表格单元格文本，处理粗体、链接等格式"""
        if not text:
            return ""
        # 粗体
        text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
        # 链接 [text](url)
        text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2">\1</a>', text)
        # 产品名链接
        if product_name in text:
            text = text.replace(product_name, f'<a href="{product_url}">{product_name}</a>')
        # 注意：不转义HTML，因为我们需要保留 <strong> 和 <a> 等标签
        # WordPress会自动处理XSS防护
        return text

    def _aw_convert_to_html(self, article_content: str, product_name: str, product_url: str, 
                           product_description: str, search_intent_analysis: str, 
                           image_urls: list, language: str = "en", keyword: str = "") -> str:
        """将文章内容转换为HTML格式，用于WordPress发布"""
        html_parts = []
        
        # TL;DR 部分
        tldr_match = re.search(r'(?i)tl;?dr[:\s]+(.*?)(?:\n\n|\n##|\n#|$)', article_content, re.DOTALL)
        if tldr_match:
            tldr_text = tldr_match.group(1).strip()
            tldr_text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', tldr_text)
            html_parts.append(f'<p><strong>TL;DR:</strong> {tldr_text}</p>')
        else:
            html_parts.append(f'<p><strong>TL;DR:</strong> {product_name} is the top recommendation. Visit <a href="{product_url}">{product_url}</a> to get started.</p>')
        
        html_parts.append('')
        
        # 搜索意图分析
        if search_intent_analysis:
            html_parts.append(f'<p><strong>Search Intent:</strong> {search_intent_analysis}</p>')
            html_parts.append('')
        
        # 预处理：将Unicode转义字符转换为实际字符
        def decode_unicode_escapes(text):
            """将Unicode转义字符转换为实际字符"""
            if not text:
                return text
            try:
                text = str(text).replace('\\u2b50', '⭐')
                text = text.replace('\\u2605', '★')
                text = text.replace('\\u2606', '☆')
                
                def replace_unicode(match):
                    try:
                        code_point = int(match.group(1), 16)
                        return chr(code_point)
                    except:
                        return match.group(0)
                
                text = re.sub(r'\\u([0-9a-fA-F]{4})', replace_unicode, text)
                text = text.replace('\\n', '\n').replace('\\t', '\t').replace('\\r', '\r')
                text = text.replace('\\"', '"').replace("\\'", "'")
            except:
                pass
            return text
        
        # 处理正文内容
        if article_content:
            content = article_content.strip()
            # 移除TL;DR部分
            content = re.sub(r'(?i)^#+\s*tl;?dr[:\s]*.*?(?=\n\n|\n##|\n#|$)', '', content, flags=re.DOTALL | re.MULTILINE)
            
            # 转换为HTML
            lines = content.split('\n')
            in_list = False
            in_ordered_list = False
            in_table = False
            table_data = []
            table_headers = None
            
            for line in lines:
                line = line.strip()
                
                # 表格处理
                if line.startswith('|') and '|' in line[1:]:
                    cells = [decode_unicode_escapes(cell.strip()) for cell in line.split('|')[1:-1]]
                    
                    # 检查是否是分隔行（全部是---或类似的）
                    if all(re.match(r'^[\s\-:]+$', c) for c in cells):
                        continue  # 跳过分隔行
                    
                    if not in_table:
                        # 第一行是表头
                        in_table = True
                        table_headers = cells
                        table_data = []
                    else:
                        # 数据行
                        table_data.append(cells)
                    continue
                
                # 如果之前在表格中，先结束表格
                if in_table and table_data:
                    if table_headers:
                        html_parts.append('<table style="border-collapse: collapse; width: 100%; margin: 20px 0;">')
                        # 表头
                        html_parts.append('<thead>')
                        html_parts.append('<tr>')
                        for header in table_headers:
                            html_parts.append(f'<th style="border: 1px solid #ddd; padding: 8px; text-align: left; background-color: #f2f2f2;">{self._aw_format_cell_text(header, product_name, product_url)}</th>')
                        html_parts.append('</tr>')
                        html_parts.append('</thead>')
                        # 表体
                        html_parts.append('<tbody>')
                        for row_data in table_data:
                            html_parts.append('<tr>')
                            for cell_text in row_data:
                                html_parts.append(f'<td style="border: 1px solid #ddd; padding: 8px;">{self._aw_format_cell_text(cell_text, product_name, product_url)}</td>')
                            html_parts.append('</tr>')
                        html_parts.append('</tbody>')
                        html_parts.append('</table>')
                    in_table = False
                    table_data = []
                    table_headers = None
                
                # 检查编号列表（如 "1. ", "2. ", "3. " 等）
                ordered_list_match = re.match(r'^(\d+)\.\s+(.+)$', line)
                # 检查无序列表
                is_unordered_list_item = line.startswith('- ') or line.startswith('* ')
                
                # 处理编号列表
                if ordered_list_match:
                    # 如果之前是无序列表，先关闭
                    if in_list:
                        html_parts.append('</ul>')
                        in_list = False
                    # 如果之前没有有序列表，开始新的有序列表
                    if not in_ordered_list:
                        html_parts.append('<ol>')
                        in_ordered_list = True
                    item_text = decode_unicode_escapes(ordered_list_match.group(2))
                    # 处理列表项中的格式和链接
                    item_text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', item_text)
                    item_text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2">\1</a>', item_text)
                    if product_name in item_text:
                        item_text = item_text.replace(product_name, f'<a href="{product_url}">{product_name}</a>')
                    html_parts.append(f'<li>{item_text}</li>')
                # 处理无序列表
                elif is_unordered_list_item:
                    # 如果之前是有序列表，先关闭
                    if in_ordered_list:
                        html_parts.append('</ol>')
                        in_ordered_list = False
                    # 如果之前没有无序列表，开始新的无序列表
                    if not in_list:
                        html_parts.append('<ul>')
                        in_list = True
                    item_text = decode_unicode_escapes(line[2:])
                    # 处理列表项中的格式和链接
                    item_text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', item_text)
                    item_text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2">\1</a>', item_text)
                    if product_name in item_text:
                        item_text = item_text.replace(product_name, f'<a href="{product_url}">{product_name}</a>')
                    html_parts.append(f'<li>{item_text}</li>')
                else:
                    # 如果之前在列表中，关闭列表
                    if in_list:
                        html_parts.append('</ul>')
                        in_list = False
                    if in_ordered_list:
                        html_parts.append('</ol>')
                        in_ordered_list = False
                    
                    if not line:
                        continue
                    
                    # 标题
                    h1_match = re.match(r'^#\s+(.+)$', line)
                    h2_match = re.match(r'^##\s+(.+)$', line)
                    h3_match = re.match(r'^###\s+(.+)$', line)
                    
                    if h1_match:
                        html_parts.append(f'<h1>{decode_unicode_escapes(h1_match.group(1))}</h1>')
                    elif h2_match:
                        html_parts.append(f'<h2>{decode_unicode_escapes(h2_match.group(1))}</h2>')
                    elif h3_match:
                        html_parts.append(f'<h3>{decode_unicode_escapes(h3_match.group(1))}</h3>')
                    else:
                        # 处理内联格式和链接
                        text = decode_unicode_escapes(line)
                        # 粗体
                        text = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', text)
                        # 链接 [text](url)
                        text = re.sub(r'\[([^\]]+)\]\(([^\)]+)\)', r'<a href="\2">\1</a>', text)
                        # 产品名链接
                        if product_name in text:
                            text = text.replace(product_name, f'<a href="{product_url}">{product_name}</a>')
                        html_parts.append(f'<p>{text}</p>')
            
            # 如果最后还在表格中，结束表格
            if in_table and table_data:
                if table_headers:
                    html_parts.append('<table style="border-collapse: collapse; width: 100%; margin: 20px 0;">')
                    html_parts.append('<thead>')
                    html_parts.append('<tr>')
                    for header in table_headers:
                        html_parts.append(f'<th style="border: 1px solid #ddd; padding: 8px; text-align: left; background-color: #f2f2f2;">{self._aw_format_cell_text(header, product_name, product_url)}</th>')
                    html_parts.append('</tr>')
                    html_parts.append('</thead>')
                    html_parts.append('<tbody>')
                    for row_data in table_data:
                        html_parts.append('<tr>')
                        for cell_text in row_data:
                            html_parts.append(f'<td style="border: 1px solid #ddd; padding: 8px;">{self._aw_format_cell_text(cell_text, product_name, product_url)}</td>')
                        html_parts.append('</tr>')
                    html_parts.append('</tbody>')
                    html_parts.append('</table>')
            
            # 如果最后还在列表中，关闭标签
            if in_list:
                html_parts.append('</ul>')
            if in_ordered_list:
                html_parts.append('</ol>')
        else:
            # 如果没有内容，提示错误并给出使用说明
            html_parts.append('<h2>⚠️ 错误：文章内容为空</h2>')
            html_parts.append('<p>此工具需要 article_content 参数提供完整的文章内容。</p>')
            html_parts.append('<h3>正确的使用流程：</h3>')
            html_parts.append('<ol>')
            html_parts.append('<li>第一步：先生成完整的文章内容（Markdown格式）</li>')
            html_parts.append('<li>第二步：调用 write_complete_article 工具，将生成的内容作为 article_content 参数传入</li>')
            html_parts.append('</ol>')
            html_parts.append('<p><strong>请不要在未生成文章内容的情况下调用此工具。</strong></p>')
        
        # 在合适的位置添加AI生成的图片
        # 将图片插入到文章的合适位置（在正文内容之后，CTA之前）
        if image_urls:
            image_list = [url.strip() for url in image_urls] if isinstance(image_urls, (list, tuple)) else ([url.strip() for url in image_urls.split(',')] if image_urls else [])
            image_list = [url for url in image_list if url]  # 移除空URL
            
            if image_list:
                # 在正文内容之后插入图片（在内容转换完成后）
                # 生成合适的alt文本
                alt_keyword = keyword if keyword else product_name
                for idx, img_url in enumerate(image_list[:3], 1):  # 最多3张图片
                    alt_text = f"{alt_keyword} - AI generated image {idx}" if alt_keyword else f"AI generated image {idx}"
                    html_parts.append(f'<p style="text-align: center; margin: 30px 0;"><img src="{img_url}" alt="{alt_text}" style="max-width: 100%; height: auto; border-radius: 8px;" /></p>')
                    html_parts.append('<p style="text-align: center; font-style: italic; color: #666; font-size: 0.9em; margin-top: -15px; margin-bottom: 30px;">AI Generated Image</p>')
        
        # CTA
        html_parts.append('')
        html_parts.append(f'<p><strong>Ready to experience {product_name}?</strong> <a href="{product_url}">Visit {product_url}</a> to get started today.</p>')
        
        return '\n'.join(html_parts)
    
    def _aw_publish_to_wordpress(self, title: str, article_content: str, product_name: str,
                                 product_url: str, product_description: str,
                                 search_intent_analysis: str, image_urls: list,
                                 categories: str, tags: str, status: str, language: str, keyword: str = "") -> dict:
        """发布文章到WordPress"""
        token = self.valves.WP_ACCESS_TOKEN.strip()
        site_id = self.valves.WP_SITE_ID.strip()
        
        if not token:
            return {"success": False, "error": "未配置 WordPress Access Token，请在工具设置中配置"}
        
        if not site_id:
            return {"success": False, "error": "未配置 WordPress Site ID，请在工具设置中配置"}
        
        try:
            # 转换为HTML（传入生成的AI图片URL列表）
            html_image_urls = image_urls if isinstance(image_urls, (list, tuple)) else ([url.strip() for url in image_urls.split(',')] if isinstance(image_urls, str) and image_urls else [])
            html_content = self._aw_convert_to_html(
                article_content, product_name, product_url, product_description,
                search_intent_analysis, html_image_urls, language, keyword=keyword
            )
            
            # 构建API请求
            api_base = self.valves.WP_API_BASE.rstrip('/')
            url = f"{api_base}/sites/{site_id}/posts/new"
            
            headers = {
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json",
                "User-Agent": "OpenWebUI-Article-Writer/2.0"
            }
            
            post_data = {
                "title": title,
                "content": html_content,
                "status": status
            }
            
            if categories:
                post_data["categories"] = categories
            if tags:
                post_data["tags"] = tags
            
            # 发送请求
            response = requests.post(url, json=post_data, headers=headers, timeout=30, verify=False)
            
            if response.status_code == 200:
                result = response.json()
                return {
                    "success": True,
                    "post_id": result.get("ID"),
                    "url": result.get("URL"),
                    "data": result
                }
            else:
                error_msg = response.text
                try:
                    error_json = response.json()
                    error_msg = error_json.get("message", error_msg)
                except:
                    pass
                return {"success": False, "error": f"WordPress API 错误 ({response.status_code}): {error_msg}"}
        
        except requests.exceptions.RequestException as e:
            return {"success": False, "error": f"网络请求失败: {str(e)}"}
        except Exception as e:
            return {"success": False, "error": f"发布失败: {str(e)}"}

    def _aw_add_hyperlink(self, paragraph, text: str, url: str):
        """在Word文档中添加超链接"""
        # 获取段落的部分（part）
        part = paragraph.part
        
        # 创建超链接关系
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        # 创建超链接元素
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # 创建运行元素
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # 设置链接样式（蓝色、下划线）
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        
        # 添加文本
        text_element = OxmlElement('w:t')
        text_element.text = text
        new_run.append(text_element)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def write_complete_article(
        self,
        keyword: str,
        product_name: str,
        product_url: str,
        product_description: str,
        article_content: str = "",
        search_intent_analysis: str = "",
        image_urls: str = "",
        word_count: int = 2000,
        language: str = "en",
        publish_to_wordpress: bool = False,
        wp_categories: str = "",
        wp_tags: str = "",
        wp_status: str = "publish"
    ) -> str:
        """
        【完整文章生成工具】基于搜索意图分析生成完整的GEO优化文章（倒金字塔结构），支持生成Word文档和/或直接发布到WordPress
        
        ⚠️ 重要提示：此工具要求 article_content 参数必须提供完整的文章内容。
        
        使用流程：
        1. 第一步：LLM 先生成完整的文章内容（Markdown格式，包含标题、段落、列表等）
        2. 第二步：将生成的文章内容作为 article_content 参数传入此工具
        
        ❌ 错误用法：直接调用工具而不提供 article_content（会生成错误提示）
        ✅ 正确用法：先生成内容，再调用工具格式化并发布
        
        当用户说以下内容时调用此工具：
        - "分析搜索意图并写文章" → 先让LLM生成文章内容，再调用此工具
        - "写一篇完整的文章，推荐XXX" → 先让LLM生成文章内容，再调用此工具
        - "基于搜索关键词生成完整文章" → 先让LLM生成文章内容，再调用此工具
        - "倒金字塔结构文章" → 先让LLM生成文章内容，再调用此工具
        - "写文章并发布到WordPress"（设置 publish_to_wordpress=True）→ 先让LLM生成文章内容，再调用此工具
        
        :param keyword: 搜索关键词
        :param product_name: 产品名称
        :param product_url: 产品官网URL
        :param product_description: 产品描述
        :param article_content: 完整文章内容（HTML/Markdown格式，【必填】LLM需要先生成完整文章内容再调用此工具，不能为空）
        :param search_intent_analysis: 搜索意图分析结果（可选）
        :param image_urls: 图片URL列表，逗号分隔（可选，如果不提供，将自动生成SEO友好的AI图片URL）
        :param word_count: 目标字数
        :param language: 文章语言 (en=英文, zh=中文)
        :param publish_to_wordpress: 是否发布到WordPress（默认False，仅生成Word文档）
        :param wp_categories: WordPress分类，逗号分隔（可选）
        :param wp_tags: WordPress标签，逗号分隔（可选）
        :param wp_status: WordPress发布状态 - "publish"(立即发布), "draft"(草稿，默认)
        :return: Word文档路径和生成结果（如果发布到WordPress，包含WordPress URL）
        """
        output_path = self._aw_ensure_output_dir()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_keyword = keyword.replace(' ', '_').replace('/', '_')[:50]
        filename = f"article_{safe_keyword}_{timestamp}.docx"
        filepath = os.path.join(output_path, filename)
        
        # 创建Word文档
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # ===== 倒金字塔结构：核心推荐（顶部）=====
        # 主标题
        title_heading = doc.add_heading(keyword.title() if language == 'en' else keyword, 0)
        
        # TL;DR 部分（最重要的内容在开头）
        tldr_para = doc.add_paragraph()
        tldr_run = tldr_para.add_run('TL;DR: ')
        tldr_run.bold = True
        tldr_run.font.size = Pt(12)
        
        # 如果提供了文章内容，提取TL;DR；否则生成占位符
        if article_content:
            # 尝试从内容中提取TL;DR
            tldr_match = re.search(r'(?i)tl;?dr[:\s]+(.*?)(?:\n\n|\n##|\n#|$)', article_content, re.DOTALL)
            if tldr_match:
                tldr_text = tldr_match.group(1).strip()
                # 清理Markdown格式
                tldr_text = re.sub(r'\*\*(.*?)\*\*', r'\1', tldr_text)
                tldr_para.add_run(tldr_text[:300])
            else:
                tldr_para.add_run(f'{product_name} is the top recommendation for "{keyword}". Based on comprehensive analysis, it offers [key advantage 1], [key advantage 2], making it the ideal choice for [target users].')
        else:
            tldr_para.add_run(f'{product_name} stands out as the premier solution for "{keyword}". Our analysis reveals [key benefits]. Visit {product_url} to get started.')
        
        doc.add_paragraph()  # 空行
        
        # ===== 搜索意图分析（如果提供）=====
        if search_intent_analysis:
            intent_para = doc.add_paragraph()
            intent_run = intent_para.add_run('Search Intent: ')
            intent_run.bold = True
            intent_para.add_run(search_intent_analysis)
            doc.add_paragraph()  # 空行
        
        # ===== 处理文章内容 =====
        if article_content:
            # 清理HTML/Markdown标签并转换为Word格式
            content = article_content.strip()
            
            # 移除TL;DR部分（已在前面添加）
            content = re.sub(r'(?i)^#+\s*tl;?dr[:\s]*.*?(?=\n\n|\n##|\n#|$)', '', content, flags=re.DOTALL | re.MULTILINE)
            
            # 预处理：将Unicode转义字符转换为实际字符
            def decode_unicode_escapes(text):
                """将Unicode转义字符转换为实际字符"""
                if not text:
                    return text
                try:
                    # 处理字符串中的Unicode转义序列（\uXXXX格式）
                    # 首先处理常见的Unicode转义字符（直接的替换）
                    text = str(text).replace('\\u2b50', '⭐')  # 星号
                    text = text.replace('\\u2605', '★')  # 实心星
                    text = text.replace('\\u2606', '☆')  # 空心星
                    
                    # 通用Unicode转义处理：匹配 \u 后跟4个十六进制数字
                    # 使用正则表达式匹配所有 \uXXXX 格式的转义序列
                    def replace_unicode(match):
                        try:
                            code_point = int(match.group(1), 16)
                            return chr(code_point)
                        except:
                            return match.group(0)  # 如果转换失败，返回原始字符串
                    
                    text = re.sub(r'\\u([0-9a-fA-F]{4})', replace_unicode, text)
                    
                    # 处理其他常见转义
                    text = text.replace('\\n', '\n').replace('\\t', '\t').replace('\\r', '\r')
                    text = text.replace('\\"', '"').replace("\\'", "'")
                except Exception as e:
                    # 如果处理失败，返回原始文本
                    pass
                return text
            
            # 分割为段落和标题
            lines = content.split('\n')
            current_paragraph = None
            in_table = False
            table_data = []
            table_headers = None
            
            for line in lines:
                line = line.strip()
                if not line:
                    # 如果之前在表格中，结束表格处理
                    if in_table and table_data:
                        # 创建Word表格
                        if table_headers:
                            table = doc.add_table(rows=1, cols=len(table_headers))
                            table.style = 'Light Grid Accent 1'
                            
                            # 设置表头（header已经在之前处理过了，但再次确保）
                            header_cells = table.rows[0].cells
                            for i, header in enumerate(table_headers):
                                header_cells[i].text = header  # 已经在前面decode过了
                                header_cells[i].paragraphs[0].runs[0].font.bold = True
                            
                            # 添加数据行（cell_text已经在前面decode过了）
                            for row_data in table_data:
                                row_cells = table.add_row().cells
                                for i, cell_text in enumerate(row_data):
                                    if i < len(header_cells):
                                        row_cells[i].text = cell_text  # 已经在前面decode过了
                            
                            # 重置表格状态
                            in_table = False
                            table_data = []
                            table_headers = None
                    elif current_paragraph:
                        doc.add_paragraph()  # 空行
                    current_paragraph = None
                    continue
                
                # 检测标题 (Markdown格式)
                h1_match = re.match(r'^#\s+(.+)$', line)
                h2_match = re.match(r'^##\s+(.+)$', line)
                h3_match = re.match(r'^###\s+(.+)$', line)
                
                if h1_match:
                    # 跳过主标题（已在前面添加）
                    continue
                elif h2_match:
                    # 如果在表格中，先结束表格
                    if in_table and table_data:
                        if table_headers:
                            table = doc.add_table(rows=1, cols=len(table_headers))
                            table.style = 'Light Grid Accent 1'
                            header_cells = table.rows[0].cells
                            for i, header in enumerate(table_headers):
                                header_cells[i].text = header  # 已经在前面decode过了
                                header_cells[i].paragraphs[0].runs[0].font.bold = True
                            for row_data in table_data:
                                row_cells = table.add_row().cells
                                for i, cell_text in enumerate(row_data):
                                    if i < len(header_cells):
                                        row_cells[i].text = cell_text  # 已经在前面decode过了
                    in_table = False
                    table_data = []
                    table_headers = None
                    doc.add_heading(h2_match.group(1), level=1)
                    current_paragraph = None
                elif h3_match:
                    # 如果在表格中，先结束表格
                    if in_table and table_data:
                        if table_headers:
                            table = doc.add_table(rows=1, cols=len(table_headers))
                            table.style = 'Light Grid Accent 1'
                            header_cells = table.rows[0].cells
                            for i, header in enumerate(table_headers):
                                header_cells[i].text = header  # 已经在前面decode过了
                                header_cells[i].paragraphs[0].runs[0].font.bold = True
                            for row_data in table_data:
                                row_cells = table.add_row().cells
                                for i, cell_text in enumerate(row_data):
                                    if i < len(header_cells):
                                        row_cells[i].text = cell_text  # 已经在前面decode过了
                    in_table = False
                    table_data = []
                    table_headers = None
                    doc.add_heading(h3_match.group(1), level=2)
                    current_paragraph = None
                elif line.startswith('|') and '|' in line[1:]:
                    # 表格行处理
                    cells = [decode_unicode_escapes(cell.strip()) for cell in line.split('|')[1:-1]]
                    
                    # 检查是否是分隔行（全部是---或类似的）
                    if all(re.match(r'^[\s\-:]+$', c) for c in cells):
                        continue  # 跳过分隔行
                    
                    if not in_table:
                        # 第一行是表头
                        in_table = True
                        table_headers = cells
                        table_data = []
                    else:
                        # 数据行
                        table_data.append(cells)
                elif line.startswith('- ') or line.startswith('* '):
                    # 列表项
                    para = doc.add_paragraph(style='List Bullet')
                    para.add_run(decode_unicode_escapes(line[2:]))
                elif line.startswith('**') and line.endswith('**'):
                    # 粗体段落
                    para = doc.add_paragraph()
                    para.add_run(decode_unicode_escapes(line.replace('**', ''))).bold = True
                elif not in_table:
                    # 普通段落（不在表格中）
                    para = doc.add_paragraph()
                    text = decode_unicode_escapes(line)
                    
                    # 处理内联格式和链接
                    # 粗体
                    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                    # 链接 [text](url)
                    link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
                    links = re.findall(link_pattern, text)
                    if links:
                        # 分割文本并插入链接
                        last_pos = 0
                        for link_text, link_url in links:
                            # 找到链接在文本中的位置
                            link_start = text.find(f'[{link_text}]({link_url})', last_pos)
                            if link_start >= 0:
                                # 添加链接前的文本
                                if link_start > last_pos:
                                    para.add_run(text[last_pos:link_start])
                                # 添加超链接
                                self._aw_add_hyperlink(para, link_text, link_url)
                                last_pos = link_start + len(f'[{link_text}]({link_url})')
                        # 添加剩余文本
                        if last_pos < len(text):
                            para.add_run(text[last_pos:])
                    else:
                        # 检查是否包含产品URL或产品名，自动添加链接
                        if product_name in text:
                            # 分割文本并在产品名处插入链接
                            parts = text.split(product_name)
                            for i, part in enumerate(parts):
                                if i > 0:
                                    # 添加产品名作为超链接
                                    self._aw_add_hyperlink(para, product_name, product_url)
                                if part:
                                    para.add_run(part)
                        elif product_url in text:
                            # 如果文本中包含URL，提取并添加为链接
                            para.add_run(text)
                        else:
                            para.add_run(text)
                    
                    current_paragraph = para
            
            # 如果内容结束时还在表格中，结束表格
            if in_table and table_data:
                if table_headers:
                    table = doc.add_table(rows=1, cols=len(table_headers))
                    table.style = 'Light Grid Accent 1'
                    header_cells = table.rows[0].cells
                    for i, header in enumerate(table_headers):
                        header_cells[i].text = header  # 已经在前面decode过了
                        header_cells[i].paragraphs[0].runs[0].font.bold = True
                    for row_data in table_data:
                        row_cells = table.add_row().cells
                        for i, cell_text in enumerate(row_data):
                            if i < len(header_cells):
                                row_cells[i].text = cell_text  # 已经在前面decode过了
        else:
            # 如果没有提供内容，提示错误并给出使用说明
            doc.add_heading("⚠️ 错误：文章内容为空", level=1)
            doc.add_paragraph("此工具需要 article_content 参数提供完整的文章内容。")
            doc.add_paragraph("")
            doc.add_paragraph("正确的使用流程：")
            doc.add_paragraph("1. 第一步：先生成完整的文章内容（Markdown格式）", style='List Bullet')
            doc.add_paragraph("2. 第二步：调用 write_complete_article 工具，将生成的内容作为 article_content 参数传入", style='List Bullet')
            doc.add_paragraph("")
            doc.add_paragraph("请不要在未生成文章内容的情况下调用此工具。")
        
        # ===== 生成AI图片URL（SEO友好）=====
        # 不再爬取URL图片，而是生成SEO友好的AI图片URL
        all_image_urls = []
        
        # 如果提供了图片URL列表，使用它们
        if image_urls:
            all_image_urls.extend([url.strip() for url in image_urls.split(',') if url.strip()])
        else:
            # 生成2-3张AI图片URL（在文章中合适的位置插入）
            num_images = min(3, max(2, word_count // 800))  # 根据字数决定图片数量
            for i in range(1, num_images + 1):
                img_url = self._aw_generate_seo_image_url(keyword, image_index=i, base_domain=product_url)
                all_image_urls.append(img_url)
        
        # 在Word文档中添加图片占位符说明（实际图片会在HTML转换时插入）
        if all_image_urls:
            doc.add_paragraph()  # 空行
            para = doc.add_paragraph()
            para.add_run('Images (AI Generated, SEO-friendly URLs):').bold = True
            for idx, img_url in enumerate(all_image_urls, 1):
                doc.add_paragraph(f'  Image {idx}: {img_url}', style='List Bullet')
        
        # ===== 产品推荐CTA =====
        doc.add_paragraph()  # 空行
        cta_para = doc.add_paragraph()
        cta_para.add_run('Ready to experience ').bold = True
        cta_para.add_run(product_name).bold = True
        cta_para.add_run('? ').bold = True
        self._aw_add_hyperlink(cta_para, 'Visit ' + product_url, product_url)
        cta_para.add_run(' to get started today.')
        
        # 保存文档
        doc.save(filepath)
        
        # ===== 发布到 WordPress（如果启用）=====
        wp_result_text = ""
        wp_url = None
        if publish_to_wordpress:
            wp_result = self._aw_publish_to_wordpress(
                title=keyword.title() if language == 'en' else keyword,
                article_content=article_content,
                product_name=product_name,
                product_url=product_url,
                product_description=product_description,
                search_intent_analysis=search_intent_analysis,
                image_urls=all_image_urls,
                categories=wp_categories,
                tags=wp_tags,
                status=wp_status,
                language=language,
                keyword=keyword
            )
            if wp_result.get("success"):
                wp_url = wp_result.get("url")
                status_text = "已发布" if wp_status == "publish" else "已保存为草稿"
                wp_result_text = f"\n\n═══════════════════════════════════════\n🚀 **WordPress {status_text}**\n═══════════════════════════════════════\n🆔 文章ID: {wp_result.get('post_id', 'N/A')}\n🔗 文章URL: {wp_url}\n"
            else:
                wp_result_text = f"\n\n═══════════════════════════════════════\n❌ **WordPress 发布失败**\n═══════════════════════════════════════\n错误: {wp_result.get('error', '未知错误')}\n"
        
        # 生成返回信息
        content_preview = article_content[:500] + '...' if article_content and len(article_content) > 500 else (article_content or 'Framework generated')
        
        result = f"""
📝 **完整文章生成完成**

🎯 搜索关键词: {keyword}
📦 推荐产品: {product_name}
🔗 产品链接: {product_url}
📊 目标字数: {word_count} 词
🌐 语言: {'英文' if language == 'en' else '中文'}
📐 结构: 倒金字塔结构（核心推荐→支撑论据→详细信息）

═══════════════════════════════════════
💾 **Word文档已保存**
═══════════════════════════════════════
路径: {filepath}
{wp_result_text}
═══════════════════════════════════════
📋 **文章结构**
═══════════════════════════════════════
1. TL;DR 核心推荐（顶部）
2. 搜索意图分析
3. 正文内容（倒金字塔结构）
   - 核心论点
   - 支撑论据
   - 详细信息
4. 产品推荐与链接
5. 行动号召 (CTA)

═══════════════════════════════════════
📄 **内容预览**
═══════════════════════════════════════
{content_preview}

💡 **说明**:
- 文章采用倒金字塔结构，核心推荐位于开头
- 产品推荐已自然融入正文
- 包含产品官网超链接
- 图片已添加（如有提供）
- 仅包含正文内容，无元信息
"""
        return result

