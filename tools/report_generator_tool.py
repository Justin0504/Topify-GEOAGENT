"""
title: 报告生成工具
description: 【项目启动报告】生成SEO+GEO项目启动报告（Word格式），包含现状分析、计划、待办事项
author: GEO Agent
version: 1.0.0
required_open_webui_version: 0.6.0
requirements: python-docx
"""

import os
from typing import List
from datetime import datetime
from pydantic import BaseModel, Field

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class Tools:
    """
    报告生成工具 - 生成专业的项目报告文档
    
    ═══════════════════════════════════════════════════════════════
    🎯 功能匹配指南（中文触发词）
    ═══════════════════════════════════════════════════════════════
    
    📄 generate_kickoff_report - 项目启动报告
       触发词: "启动报告", "项目报告", "汇报文档", "项目启动",
              "SEO报告", "GEO报告", "书面报告", "正式报告"
       示例: "生成SEO+GEO项目启动报告"
       输出: Word 文档（.docx）
    
    ═══════════════════════════════════════════════════════════════
    """

    class Valves(BaseModel):
        OUTPUT_PATH: str = Field(
            default="/app/backend/data/output",
            description="文件保存路径（Docker环境）"
        )
        COMPANY_NAME: str = Field(
            default="",
            description="公司名称（可选，用于报告署名）"
        )

    def __init__(self):
        self.valves = self.Valves()

    def _rg_ensure_output_dir(self) -> str:
        """确保输出目录存在"""
        output_path = self.valves.OUTPUT_PATH
        if not os.path.exists(output_path):
            os.makedirs(output_path, exist_ok=True)
        return output_path

    def _add_heading(self, doc, text: str, level: int = 1):
        """添加标题（doc 参数为 docx.Document 实例）"""
        heading = doc.add_heading(text, level=level)
        return heading

    def _add_paragraph(self, doc, text: str, bold: bool = False):
        """添加段落（doc 参数为 docx.Document 实例）"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        return p

    def _add_table(self, doc, headers: List[str], rows: List):
        """添加表格（doc 参数为 docx.Document 实例）"""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # 设置表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # 添加数据行
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        
        return table

    def generate_kickoff_report(
        self,
        project_name: str,
        client_name: str,
        domain: str,
        product_description: str,
        current_status: str = "",
        goals: str = "",
        timeline: str = "6个月",
        budget: str = ""
    ) -> str:
        """
        【项目启动报告工具】生成SEO+GEO项目启动报告（Word格式）
        
        当用户说以下内容时调用此工具：
        - "项目启动报告"、"生成报告"、"汇报文档"
        - "书面形式的报告"、"正式报告"
        - "SEO+GEO项目报告"
        
        :param project_name: 项目名称
        :param client_name: 客户名称
        :param domain: 网站域名
        :param product_description: 产品/服务描述
        :param current_status: 当前SEO现状描述
        :param goals: 项目目标
        :param timeline: 项目周期
        :param budget: 预算（可选）
        :return: 包含Word文件路径的结果
        """
        output_path = self._rg_ensure_output_dir()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"kickoff_report_{domain.replace('.', '_')}_{timestamp}.docx"
        filepath = os.path.join(output_path, filename)
        
        doc = Document()
        
        # ===== 封面页 =====
        doc.add_paragraph()
        doc.add_paragraph()
        
        title = doc.add_heading(f'{project_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('SEO + GEO 项目启动报告')
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        
        doc.add_paragraph()
        
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f'客户: {client_name}\n').bold = True
        info_para.add_run(f'网站: {domain}\n')
        info_para.add_run(f'日期: {datetime.now().strftime("%Y年%m月%d日")}\n')
        if self.valves.COMPANY_NAME:
            info_para.add_run(f'\n编制: {self.valves.COMPANY_NAME}')
        
        doc.add_page_break()
        
        # ===== 目录 =====
        self._add_heading(doc, '目录', 1)
        toc_items = [
            '1. 项目概述',
            '2. 现状分析',
            '3. 项目目标',
            '4. 执行策略',
            '5. 时间规划',
            '6. 待办事项',
            '7. 下一步行动'
        ]
        for item in toc_items:
            p = doc.add_paragraph(item)
            p.paragraph_format.left_indent = Inches(0.5)
        
        doc.add_page_break()
        
        # ===== 1. 项目概述 =====
        self._add_heading(doc, '1. 项目概述', 1)
        
        self._add_heading(doc, '1.1 项目背景', 2)
        doc.add_paragraph(f'''
本项目旨在为 {client_name} 提供全面的搜索引擎优化（SEO）和生成式引擎优化（GEO）服务，
帮助 {domain} 在传统搜索引擎（Google、Bing）和AI搜索引擎（ChatGPT、Perplexity、Claude、Gemini）
中获得更好的可见度和流量。
        '''.strip())
        
        self._add_heading(doc, '1.2 产品/服务介绍', 2)
        doc.add_paragraph(product_description or '[请补充产品/服务描述]')
        
        self._add_heading(doc, '1.3 项目范围', 2)
        scope_items = [
            ('项目周期', timeline),
            ('目标网站', domain),
            ('服务内容', 'SEO优化 + GEO优化'),
            ('预算', budget or '待确认')
        ]
        self._add_table(doc, ['项目', '内容'], scope_items)
        
        # ===== 2. 现状分析 =====
        self._add_heading(doc, '2. 现状分析', 1)
        
        self._add_heading(doc, '2.1 当前SEO状态', 2)
        if current_status:
            doc.add_paragraph(current_status)
        else:
            doc.add_paragraph('[待补充：基于关键词研究和技术SEO审计的结果]')
        
        self._add_heading(doc, '2.2 主要发现', 2)
        findings = [
            ('关键词覆盖', '[待分析] 当前排名关键词数量和质量'),
            ('技术问题', '[待审计] 网站技术SEO问题'),
            ('内容质量', '[待评估] 现有内容的SEO友好度'),
            ('竞争态势', '[待分析] 主要竞争对手的SEO表现'),
            ('GEO现状', '[待监测] 在AI搜索中的品牌曝光情况')
        ]
        self._add_table(doc, ['维度', '现状'], findings)
        
        # ===== 3. 项目目标 =====
        self._add_heading(doc, '3. 项目目标', 1)
        
        self._add_heading(doc, '3.1 SEO目标', 2)
        seo_goals = [
            '提升目标关键词在Google前10名的排名数量',
            '增加网站自然搜索流量',
            '提高网站技术SEO评分',
            '建立内容矩阵，覆盖更多长尾关键词'
        ]
        for goal in seo_goals:
            doc.add_paragraph(f'• {goal}', style='List Bullet')
        
        self._add_heading(doc, '3.2 GEO目标', 2)
        geo_goals = [
            '在ChatGPT、Perplexity等AI搜索中获得品牌曝光',
            '成为目标关键词AI回答的推荐品牌',
            '建立AI友好的内容结构',
            '监测并持续优化AI搜索可见度'
        ]
        for goal in geo_goals:
            doc.add_paragraph(f'• {goal}', style='List Bullet')
        
        if goals:
            self._add_heading(doc, '3.3 客户定制目标', 2)
            doc.add_paragraph(goals)
        
        # ===== 4. 执行策略 =====
        self._add_heading(doc, '4. 执行策略', 1)
        
        self._add_heading(doc, '4.1 SEO策略', 2)
        doc.add_paragraph('''
采用 Pillar-Based Marketing 内容策略，建立主题集群：
• 支柱内容（Pillar Content）：3000+字的深度指南文章
• 支撑内容（Cluster Content）：1000-1500字的细分主题文章
• 内部链接：建立支柱与支撑内容之间的链接关系
        '''.strip())
        
        self._add_heading(doc, '4.2 GEO策略', 2)
        doc.add_paragraph('''
针对AI搜索引擎优化内容：
• 内容结构：添加TL;DR摘要、FAQ、清晰的段落结构
• 可引用性：在文章开头直接回答核心问题
• 数据丰富：提供原创统计数据和案例
• 结构化数据：添加Schema.org标记
• 持续监测：跟踪目标提示词的AI回答变化
        '''.strip())
        
        # ===== 5. 时间规划 =====
        self._add_heading(doc, '5. 时间规划', 1)
        
        timeline_data = [
            ('第1-2周', '项目启动', '权限获取、现状审计、关键词研究'),
            ('第3-4周', '策略规划', '内容规划、GEO计划、技术SEO修复'),
            ('第5-8周', '内容生产', '支柱文章创作、支撑内容规划'),
            ('第9-12周', '持续优化', '内容发布、外链建设、效果监测'),
            ('第13-24周', '规模化执行', '批量内容产出、持续优化迭代')
        ]
        self._add_table(doc, ['时间', '阶段', '主要工作'], timeline_data)
        
        # ===== 6. 待办事项 =====
        self._add_heading(doc, '6. 待办事项', 1)
        
        self._add_heading(doc, '6.1 客户方待办', 2)
        client_todos = [
            ('提供Google Analytics访问权限', '高', '第1周'),
            ('提供Google Search Console访问权限', '高', '第1周'),
            ('提供网站后台访问权限（如需技术修改）', '中', '第2周'),
            ('确认目标关键词和优先级', '高', '第2周'),
            ('提供产品/服务详细资料', '中', '第1周')
        ]
        self._add_table(doc, ['事项', '优先级', '截止时间'], client_todos)
        
        self._add_heading(doc, '6.2 执行方待办', 2)
        executor_todos = [
            ('完成关键词研究报告', '高', '第2周'),
            ('完成技术SEO审计', '高', '第2周'),
            ('制定内容规划方案', '高', '第3周'),
            ('制定GEO优化计划', '高', '第3周'),
            ('开始支柱文章撰写', '高', '第4周')
        ]
        self._add_table(doc, ['事项', '优先级', '截止时间'], executor_todos)
        
        # ===== 7. 下一步行动 =====
        self._add_heading(doc, '7. 下一步行动', 1)
        
        next_steps = [
            '1. 客户确认本报告内容，提出修改意见',
            '2. 双方确认项目时间表和里程碑',
            '3. 客户提供所需访问权限',
            '4. 执行方开始关键词研究和技术审计',
            '5. 安排首次进度同步会议'
        ]
        for step in next_steps:
            doc.add_paragraph(step)
        
        # ===== 签署页 =====
        doc.add_page_break()
        self._add_heading(doc, '确认签署', 1)
        
        doc.add_paragraph()
        doc.add_paragraph('客户方确认：')
        doc.add_paragraph()
        doc.add_paragraph(f'签名：________________    日期：________________')
        doc.add_paragraph(f'姓名：                    职位：')
        
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph('执行方确认：')
        doc.add_paragraph()
        doc.add_paragraph(f'签名：________________    日期：________________')
        doc.add_paragraph(f'姓名：                    职位：')
        
        # 保存文档
        doc.save(filepath)
        
        return f"""
📄 **项目启动报告生成完成**

📁 项目: {project_name}
👤 客户: {client_name}
🌐 网站: {domain}

═══════════════════════════════════════
📋 **报告内容**
═══════════════════════════════════════
1. 项目概述 - 背景、范围、产品介绍
2. 现状分析 - SEO现状、主要发现
3. 项目目标 - SEO目标、GEO目标
4. 执行策略 - SEO策略、GEO策略
5. 时间规划 - 各阶段工作安排
6. 待办事项 - 客户方&执行方待办
7. 下一步行动 - 立即行动项
+ 确认签署页

═══════════════════════════════════════
💾 **文件已保存**
═══════════════════════════════════════
路径: {filepath}
格式: Word文档 (.docx)

💡 **使用建议**:
- 根据实际情况补充 [待分析] 部分内容
- 与客户确认后添加具体数据和指标
- 打印后可用于正式签署
"""

