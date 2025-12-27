"""
title: 文件生成工具
description: 【生成文件】支持 PDF、Word、Excel、TXT、JSON、CSV、Markdown 等多种格式，自动保存或提供下载链接
author: GEO Agent
version: 2.1.0
required_open_webui_version: 0.6.0
requirements: reportlab, python-docx, openpyxl
"""

import os
import io
import json
import csv
import base64
from typing import Dict, Any, List, Optional
from datetime import datetime
from pydantic import BaseModel, Field

# PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont

# 注册中文字体
pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))  # 宋体

# Word generation
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Excel generation
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill


class Tools:
    """
    文件生成工具 - 生成多种格式的文件
    
    ═══════════════════════════════════════════════════════════════
    🎯 快速匹配指南
    ═══════════════════════════════════════════════════════════════
    
    📄 用户说"生成 PDF"、"导出 PDF"、"保存为 PDF"
       → 调用 generate_pdf
    
    📝 用户说"生成 Word"、"导出 Word"、"保存为 Word/docx"
       → 调用 generate_word
    
    📊 用户说"生成 Excel"、"导出表格"、"保存为 Excel/xlsx"
       → 调用 generate_excel
    
    📋 用户说"生成文件"、"保存文件"、"导出文件"
       → 调用 quick_generate（自动选择格式）
    
    ═══════════════════════════════════════════════════════════════
    """

    class Valves(BaseModel):
        OUTPUT_PATH: str = Field(
            default="/app/backend/data/output",
            description="文件保存路径（Docker: /app/backend/data/output，会自动同步到本地）"
        )

    def __init__(self):
        self.valves = self.Valves()

    def _try_save_file(self, filename: str, file_bytes: bytes) -> tuple:
        """
        尝试保存文件到本地，返回 (成功?, 文件路径或错误信息)
        """
        output_dir = self.valves.OUTPUT_PATH
        try:
            os.makedirs(output_dir, exist_ok=True)
            file_path = os.path.join(output_dir, filename)
            with open(file_path, 'wb') as f:
                f.write(file_bytes)
            # 验证文件确实已保存
            if os.path.exists(file_path) and os.path.getsize(file_path) > 0:
                return True, file_path
            else:
                return False, "文件保存后验证失败"
        except Exception as e:
            return False, str(e)

    def _format_size(self, size: int) -> str:
        """格式化文件大小"""
        if size < 1024:
            return f"{size} bytes"
        elif size < 1024 * 1024:
            return f"{size / 1024:.1f} KB"
        else:
            return f"{size / (1024 * 1024):.2f} MB"

    def _get_mime_type(self, ext: str) -> str:
        """获取 MIME 类型"""
        mime_types = {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "txt": "text/plain",
            "json": "application/json",
            "csv": "text/csv",
            "md": "text/markdown",
        }
        return mime_types.get(ext.lower(), "application/octet-stream")

    def _generate_response(self, filename: str, file_bytes: bytes, file_type: str) -> str:
        """
        生成响应：先尝试保存到本地，失败则返回 Base64 下载链接
        """
        size_str = self._format_size(len(file_bytes))
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # 尝试保存到本地
        saved, result = self._try_save_file(filename, file_bytes)
        
        if saved:
            return f"""✅ 文件已保存到本地！

📄 文件名: {filename}
📁 路径: {result}
📊 大小: {size_str}
📝 类型: {file_type.upper()}
🕐 时间: {timestamp}
"""
        else:
            # 保存失败，返回 Base64 下载链接
            base64_data = base64.b64encode(file_bytes).decode("utf-8")
            mime_type = self._get_mime_type(file_type)
            data_uri = f"data:{mime_type};base64,{base64_data}"
            
            return f"""✅ 文件已生成！

📄 文件名: {filename}
📊 大小: {size_str}
📝 类型: {file_type.upper()}
🕐 时间: {timestamp}

⚠️ 无法保存到本地（{result}），请使用以下方式下载：

**📥 点击下载**（Chrome/Firefox/Edge）：
[下载 {filename}]({data_uri})

**💻 或复制以下命令到终端执行**：
```bash
echo "{base64_data}" | base64 -d > ~/Downloads/{filename}
```
"""

    def generate_pdf(
        self,
        title: str,
        content: str,
        filename: Optional[str] = None,
        author: str = "GEO Agent",
        page_size: str = "A4",
        __user__: dict = None
    ) -> str:
        """
        📄 生成 PDF 文件
        
        ✅ "生成 PDF"、"导出 PDF"、"保存为 PDF"
        
        :param title: 【必填】文档标题
        :param content: 【必填】文档内容（段落用双换行分隔）
        :param filename: 文件名（可选，自动生成）
        :param author: 作者
        :param page_size: 页面大小（A4 或 letter）
        :return: 生成结果
        """
        try:
            # 生成文件名
            if not filename:
                safe_title = "".join(c for c in title if c.isalnum() or c in ('_', '-', ' ')).strip()[:30]
                safe_title = safe_title.replace(' ', '_') or "document"
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{safe_title}_{timestamp}.pdf"
            elif not filename.endswith('.pdf'):
                filename += '.pdf'

            # 创建 PDF 到内存
            buffer = io.BytesIO()
            pagesize = A4 if page_size.upper() == "A4" else letter

            doc = SimpleDocTemplate(
                buffer,
                pagesize=pagesize,
                rightMargin=72, leftMargin=72,
                topMargin=72, bottomMargin=18,
                title=title, author=author,
            )

            styles = getSampleStyleSheet()
            
            # 使用中文字体
            chinese_font = 'STSong-Light'
            
            title_style = ParagraphStyle(
                "CustomTitle", parent=styles["Heading1"],
                fontSize=22, textColor=colors.HexColor("#2c3e50"),
                spaceAfter=30, alignment=TA_CENTER, fontName=chinese_font,
            )
            body_style = ParagraphStyle(
                "CustomBody", parent=styles["BodyText"],
                fontSize=11, leading=18, spaceAfter=12, alignment=TA_LEFT,
                fontName=chinese_font,
            )
            metadata_style = ParagraphStyle(
                "Metadata", parent=styles["Normal"],
                fontSize=9, textColor=colors.grey, alignment=TA_CENTER, spaceAfter=30,
                fontName=chinese_font,
            )
            heading1_style = ParagraphStyle(
                "ChineseHeading1", parent=styles["Heading1"],
                fontSize=16, fontName=chinese_font, spaceAfter=12,
            )
            heading2_style = ParagraphStyle(
                "ChineseHeading2", parent=styles["Heading2"],
                fontSize=14, fontName=chinese_font, spaceAfter=10,
            )
            heading3_style = ParagraphStyle(
                "ChineseHeading3", parent=styles["Heading3"],
                fontSize=12, fontName=chinese_font, spaceAfter=8,
            )

            story = []
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.2 * inch))
            metadata = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Author: {author}"
            story.append(Paragraph(metadata, metadata_style))
            story.append(Spacer(1, 0.3 * inch))

            for para in content.split("\n\n"):
                if para.strip():
                    if para.strip().startswith("# "):
                        story.append(Paragraph(para.strip()[2:], heading1_style))
                    elif para.strip().startswith("## "):
                        story.append(Paragraph(para.strip()[3:], heading2_style))
                    elif para.strip().startswith("### "):
                        story.append(Paragraph(para.strip()[4:], heading3_style))
                    else:
                        story.append(Paragraph(para.strip(), body_style))
                    story.append(Spacer(1, 0.1 * inch))

            doc.build(story)
            pdf_bytes = buffer.getvalue()
            buffer.close()

            return self._generate_response(filename, pdf_bytes, "pdf")
            
        except Exception as e:
            return f"❌ 生成 PDF 失败: {str(e)}"

    def generate_word(
        self,
        title: str,
        content: str,
        filename: Optional[str] = None,
        author: str = "GEO Agent",
        __user__: dict = None
    ) -> str:
        """
        📝 生成 Word 文件
        
        ✅ "生成 Word"、"导出 Word"、"保存为 docx"
        
        :param title: 【必填】文档标题
        :param content: 【必填】文档内容
        :param filename: 文件名（可选）
        :param author: 作者
        :return: 生成结果
        """
        try:
            if not filename:
                safe_title = "".join(c for c in title if c.isalnum() or c in ('_', '-', ' ')).strip()[:30]
                safe_title = safe_title.replace(' ', '_') or "document"
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"{safe_title}_{timestamp}.docx"
            elif not filename.endswith('.docx'):
                filename += '.docx'

            doc = Document()
            doc.core_properties.author = author
            doc.core_properties.title = title
            doc.core_properties.created = datetime.now()

            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            metadata_para = doc.add_paragraph()
            metadata_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
            metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            for para in content.split("\n\n"):
                if para.strip():
                    if para.strip().startswith("# "):
                        doc.add_heading(para.strip()[2:], level=1)
                    elif para.strip().startswith("## "):
                        doc.add_heading(para.strip()[3:], level=2)
                    elif para.strip().startswith("### "):
                        doc.add_heading(para.strip()[4:], level=3)
                    else:
                        doc.add_paragraph(para.strip())

            buffer = io.BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()

            return self._generate_response(filename, docx_bytes, "docx")
            
        except Exception as e:
            return f"❌ 生成 Word 文件失败: {str(e)}"

    def generate_excel(
        self,
        data: List[List[Any]],
        filename: Optional[str] = None,
        sheet_name: str = "Sheet1",
        headers: Optional[List[str]] = None,
        title: Optional[str] = None,
        __user__: dict = None
    ) -> str:
        """
        📊 生成 Excel 文件
        
        ✅ "生成 Excel"、"导出 Excel"、"保存为表格"
        
        :param data: 【必填】二维数据列表 [[行1], [行2], ...]
        :param filename: 文件名（可选）
        :param sheet_name: 工作表名称
        :param headers: 表头列表（可选）
        :param title: 表格标题（可选）
        :return: 生成结果
        """
        try:
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"data_{timestamp}.xlsx"
            elif not filename.endswith('.xlsx'):
                filename += '.xlsx'

            wb = Workbook()
            ws = wb.active
            ws.title = sheet_name

            current_row = 1

            if title:
                col_count = len(headers) if headers else (len(data[0]) if data else 1)
                ws.merge_cells(f"A1:{chr(64 + col_count)}1")
                title_cell = ws["A1"]
                title_cell.value = title
                title_cell.font = Font(size=16, bold=True, color="FFFFFF")
                title_cell.fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
                title_cell.alignment = Alignment(horizontal="center", vertical="center")
                ws.row_dimensions[1].height = 30
                current_row = 2

            if headers:
                for col_idx, header in enumerate(headers, start=1):
                    cell = ws.cell(row=current_row, column=col_idx)
                    cell.value = header
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill(start_color="3498db", end_color="3498db", fill_type="solid")
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                current_row += 1

            for row_data in data:
                for col_idx, value in enumerate(row_data, start=1):
                    ws.cell(row=current_row, column=col_idx, value=value)
                current_row += 1

            # 自动调整列宽
            for column in ws.columns:
                max_length = 0
                column_letter = None
                for cell in column:
                    try:
                        if hasattr(cell, "column_letter"):
                            if column_letter is None:
                                column_letter = cell.column_letter
                            if cell.value and len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                    except:
                        pass
                if column_letter:
                    ws.column_dimensions[column_letter].width = min(max_length + 2, 50)

            buffer = io.BytesIO()
            wb.save(buffer)
            xlsx_bytes = buffer.getvalue()
            buffer.close()

            return self._generate_response(filename, xlsx_bytes, "xlsx")
            
        except Exception as e:
            return f"❌ 生成 Excel 文件失败: {str(e)}"

    def generate_text(
        self,
        content: str,
        filename: Optional[str] = None,
        __user__: dict = None
    ) -> str:
        """
        📋 生成文本文件
        
        :param content: 【必填】文本内容
        :param filename: 文件名（可选）
        :return: 生成结果
        """
        try:
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"text_{timestamp}.txt"
            elif not filename.endswith('.txt'):
                filename += '.txt'

            text_bytes = content.encode('utf-8')
            return self._generate_response(filename, text_bytes, "txt")
            
        except Exception as e:
            return f"❌ 生成文本文件失败: {str(e)}"

    def generate_json(
        self,
        data: Dict[str, Any],
        filename: Optional[str] = None,
        __user__: dict = None
    ) -> str:
        """
        📋 生成 JSON 文件
        
        :param data: 【必填】字典数据
        :param filename: 文件名（可选）
        :return: 生成结果
        """
        try:
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"data_{timestamp}.json"
            elif not filename.endswith('.json'):
                filename += '.json'

            json_content = json.dumps(data, indent=2, ensure_ascii=False)
            json_bytes = json_content.encode('utf-8')
            return self._generate_response(filename, json_bytes, "json")
            
        except Exception as e:
            return f"❌ 生成 JSON 文件失败: {str(e)}"

    def generate_csv(
        self,
        data: List[List[str]],
        filename: Optional[str] = None,
        headers: Optional[List[str]] = None,
        __user__: dict = None
    ) -> str:
        """
        📊 生成 CSV 文件
        
        :param data: 【必填】二维数据列表
        :param filename: 文件名（可选）
        :param headers: 表头（可选）
        :return: 生成结果
        """
        try:
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"data_{timestamp}.csv"
            elif not filename.endswith('.csv'):
                filename += '.csv'

            output = io.StringIO()
            writer = csv.writer(output)
            if headers:
                writer.writerow(headers)
            writer.writerows(data)
            csv_bytes = output.getvalue().encode('utf-8')
            output.close()

            return self._generate_response(filename, csv_bytes, "csv")
            
        except Exception as e:
            return f"❌ 生成 CSV 文件失败: {str(e)}"

    def generate_markdown(
        self,
        content: str,
        filename: Optional[str] = None,
        __user__: dict = None
    ) -> str:
        """
        📝 生成 Markdown 文件
        
        :param content: 【必填】Markdown 内容
        :param filename: 文件名（可选）
        :return: 生成结果
        """
        try:
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"document_{timestamp}.md"
            elif not filename.endswith('.md'):
                filename += '.md'

            md_bytes = content.encode('utf-8')
            return self._generate_response(filename, md_bytes, "md")
            
        except Exception as e:
            return f"❌ 生成 Markdown 文件失败: {str(e)}"

    def quick_generate(
        self,
        content: str,
        file_type: str = "pdf",
        title: str = "Document",
        filename: Optional[str] = None,
        __user__: dict = None
    ) -> str:
        """
        🚀 快速生成文件 - 自动选择格式
        
        ✅ "生成文件"、"保存文件"、"导出文件"
        
        :param content: 【必填】文件内容
        :param file_type: 文件类型（pdf, word, txt, json, md）
        :param title: 文档标题
        :param filename: 文件名（可选）
        :return: 生成结果
        """
        file_type = file_type.lower()

        if file_type in ["pdf"]:
            return self.generate_pdf(title=title, content=content, filename=filename)
        elif file_type in ["docx", "word"]:
            return self.generate_word(title=title, content=content, filename=filename)
        elif file_type in ["txt", "text"]:
            return self.generate_text(content=content, filename=filename)
        elif file_type in ["md", "markdown"]:
            return self.generate_markdown(content=content, filename=filename)
        elif file_type in ["json"]:
            try:
                data = json.loads(content)
                return self.generate_json(data=data, filename=filename)
            except json.JSONDecodeError:
                return "❌ 内容不是有效的 JSON 格式"
        else:
            return f"❌ 不支持的文件类型: {file_type}。支持: pdf, word, txt, md, json"


# ==================== 兼容性别名 ====================
Functions = Tools
Function = Tools
