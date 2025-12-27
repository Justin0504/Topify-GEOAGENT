"""
required_open_webui_version: 0.6.0
description: Report Factory for generating Word/Excel/PDF client deliverables
requirements: python-docx, openpyxl, reportlab
"""

import base64
import io
from typing import Dict, Any, List, Optional
from datetime import datetime
from pydantic import BaseModel, Field

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class Tools:
    class Valves(BaseModel):
        default_format: str = Field(
            default="pdf",
            description="Default report format (docx, xlsx, pdf)"
        )

    def __init__(self):
        self.valves = self.Valves()

    async def generate_word_report(
        self,
        title: str,
        content: Dict[str, Any],
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Generate Word document report (.docx)
        
        :param title: Report title
        :param content: Report content dictionary with sections
        :param filename: Output filename (optional)
        :return: Base64 encoded file data
        """
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "python-docx is not installed. Please install it: pip install python-docx"
            }
        
        try:
            doc = Document()
            
            # Add title
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add date
            doc.add_paragraph(f"生成日期: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("")
            
            # Add content sections
            for section_name, section_data in content.items():
                doc.add_heading(section_name, level=1)
                
                if isinstance(section_data, list):
                    for item in section_data:
                        if isinstance(item, dict):
                            # Handle dictionary items
                            for key, value in item.items():
                                p = doc.add_paragraph(style='List Bullet')
                                p.add_run(f"{key}: ").bold = True
                                p.add_run(str(value))
                        else:
                            doc.add_paragraph(str(item), style='List Bullet')
                elif isinstance(section_data, dict):
                    # Handle nested dictionaries
                    for key, value in section_data.items():
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(f"{key}: ").bold = True
                        p.add_run(str(value))
                else:
                    doc.add_paragraph(str(section_data))
                
                doc.add_paragraph("")
            
            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Convert to base64
            file_data = base64.b64encode(buffer.read()).decode('utf-8')
            
            filename = filename or f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            
            return {
                "success": True,
                "filename": filename,
                "file_data": file_data,
                "format": "docx",
                "size_bytes": len(file_data),
                "message": f"Word 报告已生成: {filename}"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error generating Word report: {str(e)}"
            }

    async def generate_excel_report(
        self,
        data: Dict[str, List[Dict[str, Any]]],
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Generate Excel report (.xlsx)
        
        :param data: Dictionary with sheet names as keys and data rows as values
        :param filename: Output filename (optional)
        :return: Base64 encoded file data
        """
        if not EXCEL_AVAILABLE:
            return {
                "success": False,
                "error": "openpyxl is not installed. Please install it: pip install openpyxl"
            }
        
        try:
            wb = Workbook()
            # Remove default sheet
            if 'Sheet' in wb.sheetnames:
                wb.remove(wb['Sheet'])
            
            # Create sheets
            for sheet_name, rows in data.items():
                ws = wb.create_sheet(title=sheet_name[:31])  # Excel sheet name limit
                
                if not rows:
                    continue
                
                # Add headers
                headers = list(rows[0].keys())
                header_row = ws.append(headers)
                
                # Style headers
                for cell in ws[1]:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                    cell.alignment = Alignment(horizontal="center")
                
                # Add data rows
                for row in rows:
                    ws.append([row.get(h, "") for h in headers])
                
                # Auto-adjust column widths
                for column in ws.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws.column_dimensions[column_letter].width = adjusted_width
            
            # Save to bytes
            buffer = io.BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            
            # Convert to base64
            file_data = base64.b64encode(buffer.read()).decode('utf-8')
            
            filename = filename or f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            
            return {
                "success": True,
                "filename": filename,
                "file_data": file_data,
                "format": "xlsx",
                "size_bytes": len(file_data),
                "sheets": list(data.keys()),
                "message": f"Excel 报告已生成: {filename}"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error generating Excel report: {str(e)}"
            }

    async def generate_pdf_report(
        self,
        title: str,
        content: Dict[str, Any],
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Generate PDF report (.pdf)
        
        :param title: Report title
        :param content: Report content dictionary with sections
        :param filename: Output filename (optional)
        :return: Base64 encoded file data
        """
        if not PDF_AVAILABLE:
            return {
                "success": False,
                "error": "reportlab is not installed. Please install it: pip install reportlab"
            }
        
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            story = []
            
            styles = getSampleStyleSheet()
            
            # Title style
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            
            # Add title
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Add date
            date_style = styles['Normal']
            story.append(Paragraph(f"生成日期: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", date_style))
            story.append(Spacer(1, 0.3*inch))
            
            # Add content sections
            for section_name, section_data in content.items():
                # Section heading
                story.append(Paragraph(section_name, styles['Heading2']))
                story.append(Spacer(1, 0.1*inch))
                
                if isinstance(section_data, list):
                    for item in section_data:
                        if isinstance(item, dict):
                            # Create table for dictionary
                            table_data = [[Paragraph(str(k), styles['Normal']), Paragraph(str(v), styles['Normal'])] 
                                         for k, v in item.items()]
                            table = Table(table_data, colWidths=[2*inch, 4*inch])
                            table.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (0, -1), colors.grey),
                                ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, -1), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                                ('BACKGROUND', (1, 0), (1, -1), colors.beige),
                                ('GRID', (0, 0), (-1, -1), 1, colors.black)
                            ]))
                            story.append(table)
                            story.append(Spacer(1, 0.2*inch))
                        else:
                            story.append(Paragraph(f"• {str(item)}", styles['Normal']))
                            story.append(Spacer(1, 0.1*inch))
                elif isinstance(section_data, dict):
                    # Create table for dictionary
                    table_data = [[Paragraph(str(k), styles['Normal']), Paragraph(str(v), styles['Normal'])] 
                                 for k, v in section_data.items()]
                    table = Table(table_data, colWidths=[2*inch, 4*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (0, -1), colors.grey),
                        ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                        ('BACKGROUND', (1, 0), (1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 0.2*inch))
                else:
                    story.append(Paragraph(str(section_data), styles['Normal']))
                    story.append(Spacer(1, 0.2*inch))
            
            # Build PDF
            doc.build(story)
            buffer.seek(0)
            
            # Convert to base64
            file_data = base64.b64encode(buffer.read()).decode('utf-8')
            
            filename = filename or f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            
            return {
                "success": True,
                "filename": filename,
                "file_data": file_data,
                "format": "pdf",
                "size_bytes": len(file_data),
                "message": f"PDF 报告已生成: {filename}"
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error generating PDF report: {str(e)}"
            }

    async def generate_report(
        self,
        title: str,
        content: Dict[str, Any],
        format: str = "pdf",
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Generate report in specified format (docx, xlsx, or pdf)
        
        :param title: Report title
        :param content: Report content
        :param format: Output format (docx, xlsx, pdf)
        :param filename: Output filename (optional)
        :return: Base64 encoded file data
        """
        if format.lower() == "docx":
            return await self.generate_word_report(title, content, filename)
        elif format.lower() == "xlsx":
            # Convert content to Excel format
            excel_data = {}
            for section, data in content.items():
                if isinstance(data, list):
                    excel_data[section] = data
                else:
                    excel_data[section] = [{"项目": section, "内容": str(data)}]
            return await self.generate_excel_report(excel_data, filename)
        elif format.lower() == "pdf":
            return await self.generate_pdf_report(title, content, filename)
        else:
            return {
                "success": False,
                "error": f"Unsupported format: {format}. Supported formats: docx, xlsx, pdf"
            }

